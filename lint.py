"""
lint.py — 文档格式验证（中英文通用）

支持两种配置方式：
  1. dict（原有方式）：
       lint.check("report.docx", config={"tc_author": "JT", ...})

  2. DocConfig（推荐）：
       from lex_docx import DocConfig
       cfg = DocConfig(author="JT", entity_names={"forbidden": [...]})
       lint.check("report.docx", config=cfg)

扩展自定义规则：
  lint.check("doc.docx", config=cfg,
             custom_rules={"my_rule": my_fn})
  # 或通过 DocConfig.custom_lint_rules

规则列表（Phase 1 + Phase 3）：
  jt_note_format          Note 三要素（B+I+HL）
  jt_note_brackets        Note 必须有包裹标记
  no_forbidden_text       无草稿禁用文字
  no_old_project_refs     无旧项目/实体名残留
  entity_name_consistency 主体名称一致性（含 typo 检测）
  tc_author_check         TC INS/DEL author 统一
  indent_consistency      同级标题缩进一致
  defined_terms_bold      首次定义术语已加粗
  table_header_format     表格标题行有底色+加粗
  table_borders           表格有边框
  table_data_not_empty    数据行不全为空
  spelling                常见错别字
"""
from __future__ import annotations

import re
from collections import Counter
from dataclasses import dataclass, field
from io import BytesIO
from pathlib import Path
from typing import Any, Callable

from docx import Document
from docx.oxml.ns import qn

from .constants import FORBIDDEN_DRAFT_PATTERNS, HIGHLIGHT_YELLOW, JT_AUTHOR


# --------------------------------------------------------------------------- #
# LintResult                                                                   #
# --------------------------------------------------------------------------- #

@dataclass
class LintResult:
    rule: str
    passed: bool
    detail: str
    locations: list[dict[str, Any]] = field(default_factory=list)
    severity: str = "error"        # "error" | "warn" | "info"（仅 lint_cfg 模式下有意义）

    def __str__(self) -> str:
        icon  = "✅" if self.passed else "❌"
        lines = [f"{icon} [{self.severity}] {self.rule}: {self.detail}"]
        for loc in self.locations[:10]:    # 最多显示 10 条
            lines.append(f"   {loc.get('context', '')}")
        if len(self.locations) > 10:
            lines.append(f"   … 共 {len(self.locations)} 处")
        return "\n".join(lines)


# --------------------------------------------------------------------------- #
# 主入口                                                                        #
# --------------------------------------------------------------------------- #

_BUILTIN_RULES: list[str] = [
    "jt_note_format",
    "jt_note_brackets",
    "no_forbidden_text",
    "no_old_project_refs",
    "entity_name_consistency",
    "tc_author_check",
    "indent_consistency",
    "defined_terms_bold",
    "table_header_format",
    "table_borders",
    "table_data_not_empty",
    "spelling",
    "rPr_consistency",
]


def check(
    docx_path: str | Path,
    rules: list[str] | None = None,
    config: Any = None,                   # dict 或 DocConfig
    custom_rules: dict[str, Callable] | None = None,
    lint_cfg: "str | Path | dict | None" = None,  # lint config JSON 路径或已解析 dict
    profile: str | None = None,           # 显式指定 profile 名
) -> list[LintResult]:
    """
    对指定 DOCX 文件执行 lint 规则集。

    Args:
        docx_path:    DOCX 文件路径
        rules:        规则名列表，None = 全部内置规则（或 profile 中启用的规则）
        config:       dict 或 DocConfig 实例（向后兼容）
        custom_rules: {rule_name: fn(doc, config_dict, check_range)} 临时追加规则
        lint_cfg:     外部 lint config JSON 路径 / dict（Profile + Selector 模式）
                      提供时启用 severity / gate 功能，优先级高于 config
        profile:      显式指定 profile 名；None = 自动按 selectors 匹配或取第一个

    Returns:
        List[LintResult]（lint_cfg 模式下含 severity 字段）
    """
    from pathlib import Path as _Path

    # ── 统一 config 为 dict ──────────────────────────────────────────────── #
    cfg_dict = _normalize_config(config)

    # ── 合并自定义规则 ────────────────────────────────────────────────────── #
    extra_rules: dict[str, Callable] = {}
    extra_rules.update(cfg_dict.pop("custom_rules", {}) or {})
    extra_rules.update(custom_rules or {})

    # ── 规则函数映射 ─────────────────────────────────────────────────────── #
    rule_funcs: dict[str, Callable] = {
        "jt_note_format":           _check_note_format,
        "jt_note_brackets":         _check_note_brackets,
        "no_forbidden_text":        _check_no_forbidden_text,
        "no_old_project_refs":      _check_no_old_project_refs,
        "entity_name_consistency":  _check_entity_name_consistency,
        "tc_author_check":          _check_tc_author,
        "indent_consistency":       _check_indent_consistency,
        "defined_terms_bold":       _check_defined_terms_bold,
        "table_header_format":      _check_table_header_format,
        "table_borders":            _check_table_borders,
        "table_data_not_empty":     _check_table_data_not_empty,
        "spelling":                 _check_spelling,
        "rPr_consistency":          _check_rPr_consistency,
    }
    rule_funcs.update(extra_rules)

    # ── 加载文档 ──────────────────────────────────────────────────────────── #
    with open(_Path(docx_path), "rb") as f:
        doc = Document(BytesIO(f.read()))

    # ── lint_cfg 模式 ─────────────────────────────────────────────────────── #
    if lint_cfg is not None:
        return _check_with_profile(
            doc, docx_path=str(docx_path),
            base_cfg=cfg_dict, rule_funcs=rule_funcs,
            lint_cfg=lint_cfg, profile_name=profile,
            extra_rule_names=list(extra_rules.keys()),
            rules_filter=rules,
        )

    # ── 经典模式（向后兼容）──────────────────────────────────────────────── #
    if rules is None:
        rules = list(_BUILTIN_RULES) + list(extra_rules.keys())

    check_range: tuple[int, int] | None = cfg_dict.get("check_range")

    results = []
    for rule in rules:
        fn = rule_funcs.get(rule)
        if fn is None:
            results.append(LintResult(rule=rule, passed=False,
                                       detail=f"未知规则: {rule}"))
        else:
            try:
                results.append(fn(doc, cfg_dict, check_range))
            except Exception as e:
                results.append(LintResult(rule=rule, passed=False,
                                           detail=f"规则执行异常: {e}"))
    return results


def _check_with_profile(
    doc,
    docx_path: str,
    base_cfg: dict,
    rule_funcs: dict,
    lint_cfg: "str | Path | dict",
    profile_name: str | None,
    extra_rule_names: list[str],
    rules_filter: list[str] | None,
) -> list[LintResult]:
    """lint_cfg 模式：按 ResolvedProfile 运行规则，附加 severity。"""
    from . import lint_config as lc
    from pathlib import Path as _Path

    # 加载 raw cfg
    if isinstance(lint_cfg, dict):
        raw_cfg = lint_cfg
    else:
        raw_cfg = lc.load_file(lint_cfg)

    resolved = lc.resolve(raw_cfg, profile_name=profile_name, doc_path=docx_path)

    # check_range：profile > base_cfg
    check_range = resolved.check_range or base_cfg.get("check_range")

    # 确定要跑的规则集
    if rules_filter is not None:
        rule_names = rules_filter
    elif resolved.rules:
        # profile 只跑 enabled=true 的规则
        rule_names = [n for n, rc in resolved.rules.items() if rc.enabled]
    else:
        rule_names = list(_BUILTIN_RULES) + extra_rule_names

    # base_config：resolved.base_config 覆盖到 base_cfg 上
    merged_base = {**base_cfg, **resolved.base_config}

    results = []
    for rule in rule_names:
        fn = rule_funcs.get(rule)
        rc = resolved.rules.get(rule)
        severity = rc.severity if rc else "error"

        if fn is None:
            results.append(LintResult(rule=rule, passed=False,
                                       detail=f"未知规则: {rule}",
                                       severity=severity))
            continue

        # 合并 rule-level overrides
        rule_cfg = lc.apply_rule_overrides(merged_base, rule, rc.overrides if rc else {})
        rule_cfg["check_range"] = check_range

        try:
            r = fn(doc, rule_cfg, check_range)
            r.severity = severity
            results.append(r)
        except Exception as e:
            results.append(LintResult(rule=rule, passed=False,
                                       detail=f"规则执行异常: {e}",
                                       severity=severity))
    return results


# --------------------------------------------------------------------------- #
# 内部辅助                                                                      #
# --------------------------------------------------------------------------- #

def _normalize_config(config: Any) -> dict:
    """将 DocConfig 或 None 统一转为 dict。"""
    if config is None:
        return {}
    if isinstance(config, dict):
        return dict(config)
    # DocConfig
    if hasattr(config, "to_lint_config"):
        return config.to_lint_config()
    # 兜底：尝试读取常见属性
    return {
        "tc_author":    getattr(config, "tc_author", JT_AUTHOR),
        "note_prefix":  getattr(config, "note_prefix", "[JT Note: "),
        "entity_names": getattr(config, "entity_names", {}),
        "common_typos": getattr(config, "common_typos", []),
        "forbidden_draft_patterns": getattr(config, "forbidden_draft_patterns", None),
    }


def _all_paragraphs(doc):
    """生成文档所有段落，含表格单元格。返回 (para, index_or_None, loc_str)。"""
    for i, para in enumerate(doc.paragraphs):
        yield para, i, f"P{i}"
    for t_idx, table in enumerate(doc.tables):
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                for para in cell.paragraphs:
                    yield para, None, f"T{t_idx}R{r_idx}C{c_idx}"


def _para_text(para) -> str:
    return "".join(t.text or "" for t in para._element.iter(qn("w:t")))


def _run_has_b(run_el) -> bool:
    rPr = run_el.find(qn("w:rPr"))
    return rPr is not None and (
        rPr.find(qn("w:b")) is not None or
        rPr.find(qn("w:bCs")) is not None
    )


def _in_range(idx, check_range) -> bool:
    if idx is None or check_range is None:
        return True
    return check_range[0] <= idx < check_range[1]


# --------------------------------------------------------------------------- #
# Phase 1 规则                                                                  #
# --------------------------------------------------------------------------- #

def _check_note_format(doc, config, check_range) -> LintResult:
    """Note run 必须有 B + I + HL（三要素）。支持任意 note_prefix 配置。"""
    note_prefix = config.get("note_prefix", "[JT Note: ")
    detect_str  = note_prefix.lstrip("[（(").rstrip(": ").strip()  # e.g. "JT Note"
    hl_expected = config.get("note_highlight", HIGHLIGHT_YELLOW)
    failures, total = [], 0

    for para, idx, loc in _all_paragraphs(doc):
        if not _in_range(idx, check_range):
            continue
        for run_el in para._element.iter(qn("w:r")):
            text = "".join(t.text or "" for t in run_el.findall(qn("w:t")))
            if detect_str not in text:
                continue
            total += 1
            rPr    = run_el.find(qn("w:rPr"))
            has_b  = rPr is not None and rPr.find(qn("w:b")) is not None
            has_i  = rPr is not None and rPr.find(qn("w:i")) is not None
            hl_el  = rPr.find(qn("w:highlight")) if rPr is not None else None
            has_hl = hl_el is not None and hl_el.get(qn("w:val")) == hl_expected
            missing = ([k for k, v in {"bold": has_b, "italic": has_i,
                                        "highlight": has_hl}.items() if not v])
            if missing:
                failures.append({"index": idx,
                                  "context": f"{loc} 缺少: {', '.join(missing)} | {text[:40]}"})

    rule = "jt_note_format"
    if not failures:
        return LintResult(rule=rule, passed=True, detail=f"{total}/{total} OK")
    return LintResult(rule=rule, passed=False,
                      detail=f"{len(failures)}/{total} Note run 格式不完整（需 B+I+HL）",
                      locations=failures)


def _check_note_brackets(doc, config, check_range) -> LintResult:
    """Note 必须有开头标记（未被方括号包裹则报警）。"""
    note_prefix = config.get("note_prefix", "[JT Note: ")
    note_name   = config.get("note_name",   note_prefix.lstrip("[（(").rstrip(": ").strip())
    # 检测：出现 "NoteKeyword:" 但前面没有开始括号字符
    pattern = re.compile(r'(?<![(\[（])\b' + re.escape(note_name) + r'\b')
    failures = []

    for para, idx, loc in _all_paragraphs(doc):
        if not _in_range(idx, check_range):
            continue
        text = _para_text(para)
        if pattern.search(text):
            failures.append({"index": idx, "context": f"{loc}: {text[:60]}"})

    if not failures:
        return LintResult(rule="jt_note_brackets", passed=True, detail="全部有包裹标记")
    return LintResult(rule="jt_note_brackets", passed=False,
                      detail=f"{len(failures)} 处 Note 缺少开头标记",
                      locations=failures)


def _check_no_forbidden_text(doc, config, check_range) -> LintResult:
    patterns = (config.get("forbidden_draft_patterns") or FORBIDDEN_DRAFT_PATTERNS)
    failures = []
    for para, idx, loc in _all_paragraphs(doc):
        if not _in_range(idx, check_range):
            continue
        text = _para_text(para)
        for p in patterns:
            if p in text:
                failures.append({"index": idx, "context": f"{loc} 含 {p!r}: {text[:60]}"})
    if not failures:
        return LintResult(rule="no_forbidden_text", passed=True, detail="clean")
    return LintResult(rule="no_forbidden_text", passed=False,
                      detail=f"发现 {len(failures)} 处草稿禁用文字", locations=failures)


def _check_no_old_project_refs(doc, config, check_range) -> LintResult:
    forbidden = (config.get("entity_names") or {}).get("forbidden", [])
    if not forbidden:
        return LintResult(rule="no_old_project_refs", passed=True,
                          detail="未配置 forbidden 列表，跳过")
    failures = []
    for para, idx, loc in _all_paragraphs(doc):
        if not _in_range(idx, check_range):
            continue
        text = _para_text(para)
        for name in forbidden:
            if name in text:
                failures.append({"index": idx, "context": f"{loc} 含 {name!r}: {text[:60]}"})
    if not failures:
        return LintResult(rule="no_old_project_refs", passed=True, detail="clean")
    return LintResult(rule="no_old_project_refs", passed=False,
                      detail=f"发现 {len(failures)} 处旧项目/实体名残留", locations=failures)


def _check_entity_name_consistency(doc, config, check_range) -> LintResult:
    entity_cfg = config.get("entity_names") or {}
    allowed    = entity_cfg.get("allowed", [])
    forbidden  = entity_cfg.get("forbidden", [])
    typos      = config.get("common_typos", [])
    failures, stats = [], Counter()

    for para, idx, loc in _all_paragraphs(doc):
        if not _in_range(idx, check_range):
            continue
        text = _para_text(para)
        for name in allowed:
            if name in text:
                stats[name] += text.count(name)
        for name in forbidden:
            if name in text:
                failures.append({"index": idx, "context": f"{loc} 禁用 {name!r}: {text[:60]}"})
        for typo in typos:
            if typo in text:
                failures.append({"index": idx, "context": f"{loc} typo {typo!r}: {text[:60]}"})

    if not failures:
        detail = "clean"
        if allowed and stats:
            top = stats.most_common(3)
            detail = "出现: " + ", ".join(f"{n}×{c}" for n, c in top)
        return LintResult(rule="entity_name_consistency", passed=True, detail=detail)
    return LintResult(rule="entity_name_consistency", passed=False,
                      detail=f"发现 {len(failures)} 处名称问题", locations=failures)


def _check_tc_author(doc, config, check_range) -> LintResult:
    expected = config.get("tc_author", JT_AUTHOR)
    failures = []
    body = doc.element.body
    for tag in (qn("w:ins"), qn("w:del")):
        for el in body.iter(tag):
            author = el.get(qn("w:author"), "")
            if author != expected:
                tc_id = el.get(qn("w:id"), "?")
                failures.append({"index": None,
                                  "context": f"TC id={tc_id} author={author!r}（期望 {expected!r}）"})
    if not failures:
        return LintResult(rule="tc_author_check", passed=True,
                          detail=f"所有 TC author = {expected!r}")
    return LintResult(rule="tc_author_check", passed=False,
                      detail=f"{len(failures)} 个 TC author 不符",
                      locations=failures)


def _check_indent_consistency(doc, config, check_range) -> LintResult:
    paras  = doc.paragraphs
    start, end = check_range if check_range else (0, len(paras))
    style_indent: dict[str, Counter] = {}
    style_paras:  dict[str, list]   = {}

    for i in range(start, min(end, len(paras))):
        para = paras[i]
        style_name = para.style.name if para.style else "Normal"
        pPr  = para._element.find(qn("w:pPr"))
        ind  = pPr.find(qn("w:ind")) if pPr is not None else None
        left = ind.get(qn("w:left"), "0") if ind is not None else "0"
        style_indent.setdefault(style_name, Counter())[left] += 1
        style_paras.setdefault(style_name, []).append((i, left))

    failures = []
    for style_name, counter in style_indent.items():
        if len(counter) <= 1:
            continue
        majority_val, _ = counter.most_common(1)[0]
        total = sum(counter.values())
        if total < 2:
            continue
        for idx, left_val in style_paras[style_name]:
            if left_val != majority_val:
                preview = paras[idx].text[:30]
                failures.append({"index": idx,
                                  "context": (f"P{idx} [{style_name}] "
                                               f"indent={left_val}（期望 {majority_val}）: {preview}")})
    if not failures:
        return LintResult(rule="indent_consistency", passed=True, detail="同级标题缩进一致")
    return LintResult(rule="indent_consistency", passed=False,
                      detail=f"{len(failures)} 个段落缩进与同级不一致", locations=failures)


# --------------------------------------------------------------------------- #
# Phase 3 规则                                                                  #
# --------------------------------------------------------------------------- #

def _check_defined_terms_bold(doc, config, check_range) -> LintResult:
    """
    检查所有首次定义术语（中英文）是否已加粗。
    扫描逻辑与 defined_terms.auto_bold 相同，只检查不修改。
    """
    from .defined_terms import _find_term_spans, _para_full_text, _get_all_runs_with_pos

    extra_pats_raw = config.get("extra_term_patterns", [])
    extra_pats = []
    for p in extra_pats_raw:
        try:
            extra_pats.append(re.compile(p, re.IGNORECASE))
        except re.error:
            pass

    paras = doc.paragraphs
    start, end = check_range if check_range else (0, len(paras))
    failures, total = [], 0

    for i in range(start, min(end, len(paras))):
        para = paras[i]
        full = _para_full_text(para)
        spans = _find_term_spans(full, extra_pats)
        for char_start, char_end in spans:
            total += 1
            term = full[char_start:char_end]
            runs = _get_all_runs_with_pos(para)
            term_runs = [r for r, rs, re_ in runs
                         if rs >= char_start and re_ <= char_end and rs < re_]
            if not term_runs:
                continue
            if not all(_run_has_b(r) for r in term_runs):
                failures.append({"index": i,
                                  "context": f"P{i} 术语 {term!r} 未加粗: {full[:50]}"})

    if not failures:
        return LintResult(rule="defined_terms_bold", passed=True,
                          detail=f"{total}/{total} 定义术语已加粗")
    return LintResult(rule="defined_terms_bold", passed=False,
                      detail=f"{len(failures)}/{total} 术语未加粗",
                      locations=failures)


def _check_table_header_format(doc, config, check_range) -> LintResult:
    """
    检查表格标题行（默认 row 0）是否有底色或加粗。
    check_tables=[0,1,2] 可限定检查哪些表格；None = 全部。
    expected_header_shading 若指定，则精确匹配颜色值。
    """
    check_tables      = config.get("check_tables")
    expected_shading  = (config.get("expected_header_shading") or "").upper().lstrip("#")
    header_row_idx    = config.get("header_row_index", 0)
    failures = []

    for ti, table in enumerate(doc.tables):
        if check_tables is not None and ti not in check_tables:
            continue
        if len(table.rows) <= header_row_idx:
            continue
        header_row = table.rows[header_row_idx]

        for ci, cell in enumerate(header_row.cells):
            tcPr  = cell._tc.find(qn("w:tcPr"))
            shd   = tcPr.find(qn("w:shd")) if tcPr is not None else None
            fill  = (shd.get(qn("w:fill"), "") if shd is not None else "").upper()
            has_shading = fill not in ("", "AUTO", "FFFFFF")
            if expected_shading:
                has_shading = fill == expected_shading

            all_bold = True
            for para in cell.paragraphs:
                if not para.text.strip():
                    continue
                for run_el in para._element.iter(qn("w:r")):
                    if not _run_has_b(run_el):
                        all_bold = False
                        break

            if not has_shading or not all_bold:
                issues = []
                if not has_shading:
                    issues.append(f"无底色(fill={fill or '空'})")
                if not all_bold:
                    issues.append("未加粗")
                failures.append({"index": None,
                                  "context": f"T{ti}R{header_row_idx}C{ci} "
                                              f"({cell.text.strip()[:15]!r}): {', '.join(issues)}"})

    if not failures:
        return LintResult(rule="table_header_format", passed=True, detail="所有标题行格式正确")
    return LintResult(rule="table_header_format", passed=False,
                      detail=f"{len(failures)} 个标题单元格格式不符", locations=failures)


def _check_table_borders(doc, config, check_range) -> LintResult:
    """检查表格是否配置了边框（任意非 none 边框视为通过）。"""
    check_tables = config.get("check_tables")
    failures = []

    for ti, table in enumerate(doc.tables):
        if check_tables is not None and ti not in check_tables:
            continue
        tblPr   = table._tbl.find(qn("w:tblPr"))
        borders = tblPr.find(qn("w:tblBorders")) if tblPr is not None else None

        if borders is None:
            failures.append({"index": None, "context": f"T{ti}: 无 tblBorders"})
            continue

        has_real = any(
            el.get(qn("w:val"), "none") not in ("none", "nil")
            for side in ("top", "left", "bottom", "right", "insideH", "insideV")
            for el in [borders.find(qn(f"w:{side}"))]
            if el is not None
        )
        if not has_real:
            failures.append({"index": None, "context": f"T{ti}: 所有边框为 none/nil"})

    if not failures:
        return LintResult(rule="table_borders", passed=True, detail="所有表格有边框")
    return LintResult(rule="table_borders", passed=False,
                      detail=f"{len(failures)} 个表格缺少边框", locations=failures)


def _check_table_data_not_empty(doc, config, check_range) -> LintResult:
    """
    检查表格数据行（跳过标题行）是否全为空。
    min_filled_cells: 每行至少 N 个单元格有内容，默认 1。
    """
    check_tables    = config.get("check_tables")
    min_filled      = config.get("min_filled_cells", 1)
    header_row_idx  = config.get("header_row_index", 0)
    failures = []

    for ti, table in enumerate(doc.tables):
        if check_tables is not None and ti not in check_tables:
            continue
        data_rows = table.rows[header_row_idx + 1:]
        for ri, row in enumerate(data_rows, start=header_row_idx + 1):
            filled = sum(1 for c in row.cells
                         if "".join(t.text or "" for t in c._tc.iter(qn("w:t"))).strip())
            if filled < min_filled:
                preview = [c.text.strip()[:8] for c in row.cells]
                failures.append({"index": None,
                                  "context": f"T{ti}R{ri}: {filled} 格有内容（期望≥{min_filled}）: {preview}"})

    if not failures:
        return LintResult(rule="table_data_not_empty", passed=True, detail="所有数据行有内容")
    return LintResult(rule="table_data_not_empty", passed=False,
                      detail=f"{len(failures)} 行数据可能为空", locations=failures)


def _check_spelling(doc, config, check_range) -> LintResult:
    """检查 common_typos 列表中的常见错别字（中英文均支持）。"""
    typos = config.get("common_typos", [])
    if not typos:
        return LintResult(rule="spelling", passed=True, detail="未配置 common_typos，跳过")

    failures = []
    for para, idx, loc in _all_paragraphs(doc):
        if not _in_range(idx, check_range):
            continue
        text = _para_text(para)
        for typo in typos:
            if typo in text:
                failures.append({"index": idx,
                                  "context": f"{loc} {typo!r}: {text[:60]}"})

    if not failures:
        return LintResult(rule="spelling", passed=True,
                          detail=f"未发现 {len(typos)} 个常见错别字")
    return LintResult(rule="spelling", passed=False,
                      detail=f"发现 {len(failures)} 处疑似错别字", locations=failures)


def _check_rPr_consistency(doc, config, check_range) -> LintResult:
    """
    检查 TC INS run 的字体/字号是否与段落 style 一致。

    需要 config["style_rPr_map"] = {style_name: dict_or_rPr_element}
    dict 格式: {"eastAsia": "仿宋_GB2312", "sz": "24"}

    检测逻辑：
    - 遍历所有 w:ins > w:r（TC INS run）
    - 比较其 w:rPr 的字体/字号与段落 pStyle 对应的 style_rPr_map 条目
    - 不一致则报警
    """
    from copy import deepcopy

    style_rPr_map = config.get("style_rPr_map", {})
    if not style_rPr_map:
        return LintResult(rule="rPr_consistency", passed=True,
                          detail="未配置 style_rPr_map，跳过")

    def _get_rPr_attrs(rPr_el) -> dict:
        """从 rPr element 或 dict 提取关键属性用于比较。"""
        if rPr_el is None:
            return {}
        if isinstance(rPr_el, dict):
            return {k: str(v) for k, v in rPr_el.items()
                    if k in ("eastAsia", "ascii", "hAnsi", "sz", "szCs")}
        attrs = {}
        rFonts = rPr_el.find(qn("w:rFonts"))
        if rFonts is not None:
            for attr_name, qname in [
                ("eastAsia", qn("w:eastAsia")),
                ("ascii", qn("w:ascii")),
                ("hAnsi", qn("w:hAnsi")),
            ]:
                val = rFonts.get(qname)
                if val:
                    attrs[attr_name] = val
        for tag_name, key in [("w:sz", "sz"), ("w:szCs", "szCs")]:
            el = rPr_el.find(qn(tag_name))
            if el is not None:
                attrs[key] = el.get(qn("w:val"), "")
        return attrs

    paras = doc.paragraphs
    start, end = check_range if check_range else (0, len(paras))
    failures = []

    for i in range(start, min(end, len(paras))):
        para = paras[i]
        para_el = para._element

        # 获取段落 style
        pPr = para_el.find(qn("w:pPr"))
        pStyle_el = pPr.find(qn("w:pStyle")) if pPr is not None else None
        style_name = pStyle_el.get(qn("w:val"), "Normal") if pStyle_el is not None else "Normal"

        expected_entry = style_rPr_map.get(style_name)
        if expected_entry is None:
            continue
        expected_attrs = _get_rPr_attrs(expected_entry)
        if not expected_attrs:
            continue

        # 检查每个 INS run
        for ins_el in para_el.findall(qn("w:ins")):
            for r_el in ins_el.findall(qn("w:r")):
                run_rPr = r_el.find(qn("w:rPr"))
                actual_attrs = _get_rPr_attrs(run_rPr)
                mismatches = []
                for key, exp_val in expected_attrs.items():
                    act_val = actual_attrs.get(key)
                    if act_val and act_val != exp_val:
                        mismatches.append(f"{key}: 实际={act_val!r} 期望={exp_val!r}")
                if mismatches:
                    text_preview = "".join(
                        t.text or "" for t in r_el.findall(qn("w:t"))
                    )[:30]
                    failures.append({
                        "index": i,
                        "context": (f"P{i} [{style_name}] INS run {text_preview!r}: "
                                    f"{'; '.join(mismatches)}")
                    })

    if not failures:
        return LintResult(rule="rPr_consistency", passed=True,
                          detail="所有 INS run 字体/字号与 style 一致")
    return LintResult(rule="rPr_consistency", passed=False,
                      detail=f"{len(failures)} 个 INS run 格式与 style 不符",
                      locations=failures)
