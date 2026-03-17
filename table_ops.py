"""
table_ops.py — 表格操作

extract_table  从 DOCX（如 AutoDocs）按 near_text 或索引提取表格数据
fill_table     按列映射填充目标表格，TC INS 形式
fill_kv_table  填充 KV 表（字段名—值 两列，如基本信息表）
adjust_rows    增删数据行（行级 TC INS / TC DEL）
format_table   统一表格格式（标题行底色、加粗、边框、列宽、对齐）
"""
from __future__ import annotations

from copy import deepcopy
from io import BytesIO
from pathlib import Path
from typing import Any

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from .constants import (
    DEFAULT_BORDER_COLOR, DEFAULT_BORDER_STYLE, DEFAULT_BORDER_WIDTH,
    HEADER_SHADING_HEX, JT_AUTHOR,
)
from .tc_utils import (
    _utc_now, make_run, make_tc_tag, mark_row_as_deleted,
    mark_row_as_inserted, next_tc_id,
)


def _get_cfg_attr(cfg, attr: str, fallback):
    if cfg is None:
        return fallback
    return getattr(cfg, attr, fallback)


# =========================================================================== #
# extract_table                                                                #
# =========================================================================== #

def extract_table(
    docx_path: str | Path,
    near_text: str | None = None,
    table_index: int | None = None,
    output: str = "list_of_dicts",
    search_range: tuple[int, int] | None = None,
) -> Any:
    """
    从 DOCX 文件中提取表格数据。

    Args:
        docx_path:    源文件路径（通常是 AutoDocs 输出文件）
        near_text:    定位关键字：取文档中第一个包含此文字的段落之后的表格
        table_index:  直接按索引取表格（与 near_text 二选一）
        output:       "list_of_dicts"（默认）| "list_of_lists" | "raw"
        search_range: 限制 near_text 的搜索范围 (para_start, para_end)

    Returns:
        list_of_dicts: [{"列名": "值", ...}, ...]  —— 首行为列名
        list_of_lists: [[cell, ...], ...]           —— 含首行
        raw:           python-docx Table 对象
    """
    with open(Path(docx_path), "rb") as f:
        doc = Document(BytesIO(f.read()))

    if table_index is not None:
        table = doc.tables[table_index]
    elif near_text is not None:
        table = _find_table_near_text(doc, near_text, search_range)
        if table is None:
            raise ValueError(f"未找到 near_text={near_text!r} 附近的表格")
    else:
        raise ValueError("必须指定 near_text 或 table_index")

    return _table_to_output(table, output)


def _find_table_near_text(doc, near_text: str, search_range) -> Any:
    """扫描 body，找第一个包含 near_text 的段落，返回其后紧跟的表格。"""
    body = doc.element.body
    para_idx = 0
    start = search_range[0] if search_range else 0
    end   = search_range[1] if search_range else None

    found_para = False
    for child in body:
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if tag == "p":
            if end is not None and para_idx >= end:
                return None
            if para_idx >= start:
                text = "".join(t.text or "" for t in child.iter(qn("w:t")))
                if near_text in text:
                    found_para = True
            para_idx += 1
        elif tag == "tbl" and found_para:
            return _tbl_el_to_table(child, doc)
    return None


def _tbl_el_to_table(tbl_el, doc):
    """将 lxml w:tbl 元素转换为 python-docx Table 对象。"""
    for t in doc.tables:
        if t._tbl is tbl_el:
            return t
    return None


def _cell_text(cell) -> str:
    """提取单元格全文（包含 w:ins 内容）。"""
    return "".join(t.text or "" for t in cell._tc.iter(qn("w:t"))).strip()


def _table_to_output(table, output_format: str) -> Any:
    if not table.rows:
        return []
    if output_format == "raw":
        return table
    if output_format == "list_of_lists":
        return [[_cell_text(c) for c in row.cells] for row in table.rows]
    # list_of_dicts（默认）
    headers = [_cell_text(c) for c in table.rows[0].cells]
    # 去掉空列头（merged cells 产生的重复）
    seen = set()
    clean_headers = []
    for h in headers:
        if h in seen:
            clean_headers.append(f"_col{len(clean_headers)}")
        else:
            clean_headers.append(h)
            seen.add(h)
    result = []
    for row in table.rows[1:]:
        d = {clean_headers[i]: _cell_text(c)
             for i, c in enumerate(row.cells)
             if i < len(clean_headers)}
        result.append(d)
    return result


# =========================================================================== #
# fill_table                                                                   #
# =========================================================================== #

def fill_table(
    doc,
    table_index: int,
    data: list[dict],
    column_mapping: dict[str, str] | None = None,
    as_tc_ins: bool = True,
    author: str | None = None,
    header_row_index: int = 0,
    auto_add_rows: bool = False,
    preserve_rPr: bool = False,
    cfg=None,
) -> int:
    """
    按列映射将 data 写入目标表格，跳过 header row。

    column_mapping:
      {源列名: 目标列名}       — 将源数据列映射到目标表格对应列
      {源列名: "auto_number"}  — 该列改为自动生成 1, 2, 3...

    Args:
        header_row_index: 标题行索引（默认 0），数据行从 header_row_index+1 开始
        auto_add_rows:    True = 数据行不足时自动调用 adjust_rows 增行（TC INS）
        preserve_rPr:     True = 写入时保留单元格原有 run 的字体/字号等 rPr
        cfg:              DocConfig（提供 author 等默认值）

    Returns:
        实际填充的行数
    """
    author = author or _get_cfg_attr(cfg, "author", JT_AUTHOR)
    table = doc.tables[table_index]
    if not table.rows:
        return 0

    # 自动增行
    if auto_add_rows and data:
        current_data_rows = len(table.rows) - header_row_index - 1
        if len(data) > current_data_rows:
            adjust_rows(doc, table_index, len(data),
                        author=author, header_row_index=header_row_index, cfg=cfg)

    # 构建目标列名 → 列索引
    header_cells = [_cell_text(c) for c in table.rows[header_row_index].cells]
    col_index: dict[str, int] = {}
    for i, h in enumerate(header_cells):
        if h and h not in col_index:
            col_index[h] = i
        # 模糊备用
        col_index.setdefault(h.strip(), i)

    # 解析 column_mapping → (src_key, action, tgt_col_idx)
    mapped: list[tuple[str, str, int | None]] = []
    if column_mapping:
        for src_key, tgt_val in column_mapping.items():
            if tgt_val == "auto_number":
                # 找目标表头中与 src_key 最相近的列
                tgt_idx = _fuzzy_col_index(header_cells, src_key)
                mapped.append((src_key, "auto", tgt_idx))
            else:
                tgt_idx = col_index.get(tgt_val, _fuzzy_col_index(header_cells, tgt_val))
                mapped.append((src_key, "col", tgt_idx))
    else:
        # 无映射：按源数据列名直接匹配目标列名
        for src_key in (data[0].keys() if data else []):
            tgt_idx = col_index.get(src_key, _fuzzy_col_index(header_cells, src_key))
            mapped.append((src_key, "col", tgt_idx))

    data_rows = table.rows[header_row_index + 1:]
    tc_id = next_tc_id(doc)
    date  = _utc_now()
    filled = 0

    for row_idx, row_data in enumerate(data):
        if row_idx >= len(data_rows):
            break
        row = data_rows[row_idx]
        for src_key, action, tgt_idx in mapped:
            if tgt_idx is None or tgt_idx >= len(row.cells):
                continue
            value = str(row_idx + 1) if action == "auto" else str(row_data.get(src_key, ""))
            _fill_cell(row.cells[tgt_idx], value, tc_id, author, date, as_tc_ins, preserve_rPr)
            tc_id += 1
        filled += 1

    return filled


def _normalize_key(s: str) -> str:
    """归一化列名：压缩所有空白（含换行）为单空格，strip。"""
    import re
    return re.sub(r'\s+', ' ', s).strip()


def _fuzzy_col_index(headers: list[str], target: str) -> int | None:
    """
    模糊匹配：找包含 target 的列头索引，找不到返回 None。
    归一化空白后匹配，避免 w:t 不含换行导致的不匹配。
    """
    t_norm = _normalize_key(target)
    for i, h in enumerate(headers):
        h_norm = _normalize_key(h)
        if t_norm == h_norm or t_norm in h_norm or h_norm in t_norm:
            return i
    return None


# =========================================================================== #
# fill_kv_table                                                                #
# =========================================================================== #

def fill_kv_table(
    doc,
    table_index: int,
    data: dict[str, str],
    key_column: int | None = None,
    value_column: int | None = None,
    as_tc_ins: bool = True,
    author: str | None = None,
    fuzzy_key: bool = True,
    key_columns: list[int] | None = None,
    preserve_rPr: bool = False,
    cfg=None,
) -> int:
    """
    填充 KV 表（字段名—值 两列格式，如基本信息表）。

    支持场景：
    - 标准两列：key_column=0, value_column=1
    - 四列布局（每行两组 KV）：key_columns=[0, 2]，value 自动取 key_col+1
    - compound key：data key 含 "/" 时拆分后逐一尝试匹配
      例：{"法定代表人/负责人/执行事务合伙人": "张三"} → 匹配文档中任意一种写法

    Args:
        key_column:   单列模式下 key 所在列（默认 0）
        value_column: 单列模式下 value 所在列（默认 key_column+1）
        key_columns:  多列模式，指定多个 key 列；value 列自动为各 key 列 +1
                      设置此参数时 key_column/value_column 被忽略
        fuzzy_key:    True = 允许 key 文字互含；False = 精确匹配

    Returns:
        实际填充的单元格数
    """
    author = author or _get_cfg_attr(cfg, "author", JT_AUTHOR)
    table = doc.tables[table_index]
    tc_id = next_tc_id(doc)
    date  = _utc_now()
    filled = 0

    # 确定 (key_col, val_col) 对列表
    if key_columns is not None:
        col_pairs = [(kc, kc + 1) for kc in key_columns]
    else:
        kc = key_column if key_column is not None else 0
        vc = value_column if value_column is not None else kc + 1
        col_pairs = [(kc, vc)]

    # 预处理 data：展开 compound key，建立 {normalized_key: (original_key, value)}
    # compound key 用 "/" 分隔，每个片段都能独立匹配
    expanded: list[tuple[str, str]] = []   # [(norm_key_fragment, value), ...]
    for raw_key, val in data.items():
        for fragment in raw_key.split("/"):
            norm = _normalize_key(fragment)
            if norm:
                expanded.append((norm, str(val)))

    for row in table.rows:
        n_cells = len(row.cells)
        for key_col, val_col in col_pairs:
            if key_col >= n_cells or val_col >= n_cells:
                continue
            cell_key_norm = _normalize_key(_cell_text(row.cells[key_col]))
            if not cell_key_norm:
                continue
            for frag_norm, val in expanded:
                if not val:
                    continue
                if fuzzy_key:
                    hit = (frag_norm == cell_key_norm or
                           frag_norm in cell_key_norm or
                           cell_key_norm in frag_norm)
                else:
                    hit = frag_norm == cell_key_norm
                if hit:
                    _fill_cell(row.cells[val_col], val, tc_id, author, date, as_tc_ins, preserve_rPr)
                    tc_id += 1
                    filled += 1
                    break   # 该 (row, col_pair) 已匹配，不继续尝试其他 fragment

    return filled


# =========================================================================== #
# adjust_rows                                                                  #
# =========================================================================== #

def adjust_rows(
    doc,
    table_index: int,
    target_data_rows: int,
    as_tc_ins: bool = True,
    author: str | None = None,
    inherit_format_from: str | int = "row_1",
    header_row_index: int = 0,
    cfg=None,
) -> dict:
    """
    调整表格数据行数到 target_data_rows（不含标题行）。

    - 行数不足：deepcopy 模板行，以行级 TC INS 标记插入
    - 行数过多：以行级 TC DEL 标记多余行

    Args:
        inherit_format_from: "row_1"（第一数据行）| "last"（最后数据行）| 整数行索引

    Returns:
        {"added": n, "deleted": n}
    """
    author = author or _get_cfg_attr(cfg, "author", JT_AUTHOR)
    table = doc.tables[table_index]
    data_rows = table.rows[header_row_index + 1:]   # 不含 header
    current = len(data_rows)
    tc_id = next_tc_id(doc)
    date  = _utc_now()
    added = deleted = 0

    # 确定模板行
    if isinstance(inherit_format_from, int):
        tpl_idx = inherit_format_from
    elif inherit_format_from == "last":
        tpl_idx = len(table.rows) - 1
    else:  # "row_1"
        tpl_idx = 1 if len(table.rows) > 1 else 0
    template_tr = table.rows[tpl_idx]._tr

    if target_data_rows > current:
        # 增行
        for _ in range(target_data_rows - current):
            new_tr = deepcopy(template_tr)
            # 清空单元格文字（保留格式）
            for tc_el in new_tr.findall(qn("w:tc")):
                for p_el in tc_el.findall(qn("w:p")):
                    pPr = p_el.find(qn("w:pPr"))
                    for child in list(p_el):
                        if child is not pPr:
                            p_el.remove(child)
            if as_tc_ins:
                mark_row_as_inserted(new_tr, tc_id, author, date)
                tc_id += 1
            table._tbl.append(new_tr)
            added += 1

    elif target_data_rows < current:
        # 删行（从最后一行往前标记 DEL）
        rows_to_del = list(data_rows)[target_data_rows:]
        for row in rows_to_del:
            if as_tc_ins:
                mark_row_as_deleted(row._tr, tc_id, author, date)
                tc_id += 1
            deleted += 1

    return {"added": added, "deleted": deleted}


# =========================================================================== #
# format_table                                                                 #
# =========================================================================== #

def format_table(
    doc,
    table_index: int,
    header_shading: str | None = None,
    header_bold: bool = True,
    borders: str | None = DEFAULT_BORDER_STYLE,
    border_width: int = DEFAULT_BORDER_WIDTH,
    border_color: str = DEFAULT_BORDER_COLOR,
    column_widths: list[int] | None = None,
    column_alignments: list[str] | None = None,
    header_row_index: int = 0,
    cfg=None,
) -> None:
    """
    统一表格格式。

    Args:
        header_shading:     标题行底色十六进制（不含 #），如 "D9E2F3"；None = 取 cfg 或默认值
        header_bold:        标题行文字加粗
        borders:            边框样式 "single" | "none" | None（不修改）
        border_width:       边框宽度（单位 1/8 磅，4 = 半磅）
        border_color:       边框颜色十六进制
        column_widths:      各列宽度（单位 dxa = 1/20 磅），如 [800, 4000, 2000]
        column_alignments:  各列对齐，"left" | "center" | "right"
        header_row_index:   标题行位置（默认 0）
        cfg:                DocConfig（提供 header_shading / border_* 等默认值）
    """
    header_shading = header_shading or _get_cfg_attr(cfg, "header_shading", HEADER_SHADING_HEX)
    borders        = borders        or _get_cfg_attr(cfg, "border_style",   DEFAULT_BORDER_STYLE)
    border_width   = border_width   or _get_cfg_attr(cfg, "border_width",   DEFAULT_BORDER_WIDTH)
    border_color   = border_color   or _get_cfg_attr(cfg, "border_color",   DEFAULT_BORDER_COLOR)

    table = doc.tables[table_index]
    if not table.rows:
        return

    # ── 标题行格式 ─────────────────────────────────────────────────────────── #
    if header_shading or header_bold:
        for cell in table.rows[header_row_index].cells:
            if header_shading:
                _set_cell_shading(cell, header_shading)
            if header_bold:
                _bold_cell_runs(cell)

    # ── 全表边框 ──────────────────────────────────────────────────────────── #
    if borders is not None:
        _set_table_borders(table, borders, border_width, border_color)

    # ── 列宽 ──────────────────────────────────────────────────────────────── #
    if column_widths:
        _set_column_widths(table, column_widths)

    # ── 列对齐 ────────────────────────────────────────────────────────────── #
    if column_alignments:
        _set_column_alignments(table, column_alignments)


# =========================================================================== #
# 内部辅助                                                                     #
# =========================================================================== #

def _fill_cell(cell, text: str, tc_id: int, author: str, date: str,
               as_tc_ins: bool = True, preserve_rPr: bool = False) -> None:
    """
    向单元格首段写入文字，清空原有 run（保留 pPr）。
    as_tc_ins=True 时用 w:ins 包裹。
    preserve_rPr=True 时，写入前先保存首个 run 的 rPr，写入后应用到新 run。
    """
    if not cell.paragraphs:
        p_el = OxmlElement("w:p")
        cell._tc.append(p_el)
    else:
        p_el = cell.paragraphs[0]._element

    # 保留原有 rPr
    saved_rPr = None
    if preserve_rPr:
        first_run = p_el.find(qn("w:r"))
        if first_run is not None:
            rPr_el = first_run.find(qn("w:rPr"))
            if rPr_el is not None:
                saved_rPr = deepcopy(rPr_el)

    # 清空非 pPr 子元素
    pPr = p_el.find(qn("w:pPr"))
    for child in list(p_el):
        if child is not pPr:
            p_el.remove(child)

    new_run = make_run(text)
    if saved_rPr is not None:
        new_run.insert(0, saved_rPr)

    if as_tc_ins:
        ins = make_tc_tag("w:ins", tc_id, author, date)
        ins.append(new_run)
        p_el.append(ins)
    else:
        p_el.append(new_run)


def _set_cell_shading(cell, hex_color: str) -> None:
    """设置单元格底色（w:tcPr > w:shd）。"""
    tcPr = cell._tc.find(qn("w:tcPr"))
    if tcPr is None:
        tcPr = OxmlElement("w:tcPr")
        cell._tc.insert(0, tcPr)
    existing = tcPr.find(qn("w:shd"))
    if existing is not None:
        tcPr.remove(existing)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color.lstrip("#").upper())
    tcPr.append(shd)


def _bold_cell_runs(cell) -> None:
    """对单元格内所有 run 加粗（w:b + w:bCs）。"""
    for para in cell.paragraphs:
        for run_el in para._element.iter(qn("w:r")):
            rPr = run_el.find(qn("w:rPr"))
            if rPr is None:
                rPr = OxmlElement("w:rPr")
                run_el.insert(0, rPr)
            if rPr.find(qn("w:b")) is None:
                rPr.append(OxmlElement("w:b"))
            if rPr.find(qn("w:bCs")) is None:
                rPr.append(OxmlElement("w:bCs"))


def _set_table_borders(table, style: str, width: int, color: str) -> None:
    """设置全表边框（w:tblPr > w:tblBorders）。"""
    tblPr = table._tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        table._tbl.insert(0, tblPr)
    existing = tblPr.find(qn("w:tblBorders"))
    if existing is not None:
        tblPr.remove(existing)

    tblBorders = OxmlElement("w:tblBorders")
    hex_color = color.lstrip("#").upper()
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), style)
        el.set(qn("w:sz"), str(width))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), hex_color)
        tblBorders.append(el)
    tblPr.append(tblBorders)


def _set_column_widths(table, widths: list[int]) -> None:
    """
    设置列宽：更新 w:tblGrid + 每行每列的 w:tcPr > w:tcW。
    widths 单位为 dxa（twentieths of a point）。
    """
    tbl = table._tbl

    # 更新 w:tblGrid
    existing_grid = tbl.find(qn("w:tblGrid"))
    if existing_grid is not None:
        tbl.remove(existing_grid)
    tblGrid = OxmlElement("w:tblGrid")
    for w in widths:
        gridCol = OxmlElement("w:gridCol")
        gridCol.set(qn("w:w"), str(w))
        tblGrid.append(gridCol)
    # 插在 tblPr 之后
    tblPr = tbl.find(qn("w:tblPr"))
    insert_pos = list(tbl).index(tblPr) + 1 if tblPr is not None else 0
    tbl.insert(insert_pos, tblGrid)

    # 更新每行每格的 tcW
    for row in table.rows:
        for ci, cell in enumerate(row.cells):
            if ci >= len(widths):
                break
            tcPr = cell._tc.find(qn("w:tcPr"))
            if tcPr is None:
                tcPr = OxmlElement("w:tcPr")
                cell._tc.insert(0, tcPr)
            existing_tcW = tcPr.find(qn("w:tcW"))
            if existing_tcW is not None:
                tcPr.remove(existing_tcW)
            tcW = OxmlElement("w:tcW")
            tcW.set(qn("w:w"), str(widths[ci]))
            tcW.set(qn("w:type"), "dxa")
            tcPr.append(tcW)


def _set_column_alignments(table, alignments: list[str]) -> None:
    """
    设置列对齐：对每列的所有段落设置 w:pPr > w:jc。
    alignments: ["left", "center", "right", ...]
    Word 值：left / center / right / both（两端对齐）
    """
    _jc_map = {"left": "left", "center": "center", "right": "right",
                "justify": "both", "both": "both"}
    for row in table.rows:
        for ci, cell in enumerate(row.cells):
            if ci >= len(alignments):
                break
            jc_val = _jc_map.get(alignments[ci], alignments[ci])
            for para in cell.paragraphs:
                pPr = para._element.get_or_add_pPr()
                existing_jc = pPr.find(qn("w:jc"))
                if existing_jc is not None:
                    pPr.remove(existing_jc)
                jc = OxmlElement("w:jc")
                jc.set(qn("w:val"), jc_val)
                pPr.append(jc)


# =========================================================================== #
# copy_table  (FR-9)                                                           #
# =========================================================================== #

def copy_table(
    src_doc,
    dst_doc,
    dst_position: str,
    src_table_index: int | None = None,
    src_near_text: str | None = None,
    transform: dict | None = None,
    as_tc_ins: bool = True,
    author: str | None = None,
    cfg=None,
) -> None:
    """
    跨文档复制表格（含完整格式），插入目标文档指定位置。

    Args:
        src_doc:          源文档路径（str/Path）或已加载的 Document 对象
        dst_doc:          目标 Document 对象
        dst_position:     插入/替换位置：
                            "after_para:241"    — 在第 241 个段落之后插入
                            "replace_table:15"  — TC DEL 第 15 个表格，再插入新表
        src_table_index:  源表格索引（与 src_near_text 二选一）
        src_near_text:    源文档中定位用的临近文字
        transform:        可选变换（见下）
        as_tc_ins:        True = 新表格每行以 TC INS 标记
        author:           TC 作者
        cfg:              DocConfig

    transform 字典支持的键：
        columns:        list[int]  — 只保留这些列索引（0-based）
        rename_headers: dict       — {旧列头: 新列头}，修改标题行文字
        max_rows:       int        — 数据行上限（不含标题行）
        filter_rows:    callable   — fn(row_dict) -> bool，返回 False 则跳过该行
    """
    from io import BytesIO
    from pathlib import Path as _Path

    author = author or _get_cfg_attr(cfg, "author", JT_AUTHOR)

    # ── 加载源文档 ──────────────────────────────────────────────────────────── #
    if isinstance(src_doc, (str, _Path)):
        from docx import Document as _Doc
        with open(_Path(src_doc), "rb") as f:
            _src = _Doc(BytesIO(f.read()))
    else:
        _src = src_doc

    # ── 定位源表格 ─────────────────────────────────────────────────────────── #
    if src_table_index is not None:
        src_tbl = _src.tables[src_table_index]
    elif src_near_text is not None:
        src_tbl = _find_table_near_text(_src, src_near_text, None)
        if src_tbl is None:
            raise ValueError(f"未找到 src_near_text={src_near_text!r} 附近的表格")
    else:
        raise ValueError("必须指定 src_table_index 或 src_near_text")

    # ── deepcopy 源表格 XML ────────────────────────────────────────────────── #
    new_tbl = deepcopy(src_tbl._tbl)

    # ── 应用 transform ────────────────────────────────────────────────────── #
    if transform:
        new_tbl = _apply_transform(new_tbl, src_tbl, transform)

    # ── TC INS 标记每行 ───────────────────────────────────────────────────── #
    if as_tc_ins:
        tc_id = next_tc_id(dst_doc)
        date  = _utc_now()
        for tr in new_tbl.findall(qn("w:tr")):
            mark_row_as_inserted(tr, tc_id, author, date)
            tc_id += 1

    # ── 解析目标位置并插入 ────────────────────────────────────────────────── #
    _insert_table_at(dst_doc, new_tbl, dst_position, author, as_tc_ins)


def copy_table_format(
    src_doc,
    dst_doc,
    src_table_index: int,
    dst_table_index: int,
) -> None:
    """
    只复制格式（不复制数据）：从源表格复制 tblPr / tblGrid / 列宽 / 单元格格式
    到目标表格，保留目标表格的单元格内容。

    Args:
        src_doc:          源文档路径或 Document 对象
        dst_doc:          目标 Document 对象
        src_table_index:  源表格索引
        dst_table_index:  目标表格索引
    """
    from io import BytesIO
    from pathlib import Path as _Path

    if isinstance(src_doc, (str, _Path)):
        from docx import Document as _Doc
        with open(_Path(src_doc), "rb") as f:
            _src = _Doc(BytesIO(f.read()))
    else:
        _src = src_doc

    src_tbl_el = _src.tables[src_table_index]._tbl
    dst_tbl_el = dst_doc.tables[dst_table_index]._tbl

    # ── tblPr（表格属性：边框/底色/布局）──────────────────────────────────── #
    src_tblPr = src_tbl_el.find(qn("w:tblPr"))
    if src_tblPr is not None:
        existing = dst_tbl_el.find(qn("w:tblPr"))
        if existing is not None:
            dst_tbl_el.remove(existing)
        dst_tbl_el.insert(0, deepcopy(src_tblPr))

    # ── tblGrid（列宽定义）────────────────────────────────────────────────── #
    src_grid = src_tbl_el.find(qn("w:tblGrid"))
    if src_grid is not None:
        existing = dst_tbl_el.find(qn("w:tblGrid"))
        if existing is not None:
            dst_tbl_el.remove(existing)
        tblPr = dst_tbl_el.find(qn("w:tblPr"))
        insert_pos = list(dst_tbl_el).index(tblPr) + 1 if tblPr is not None else 0
        dst_tbl_el.insert(insert_pos, deepcopy(src_grid))

    # ── 行/单元格属性（逐行逐格复制 trPr / tcPr，不动 w:p 内容）────────────── #
    src_rows = src_tbl_el.findall(qn("w:tr"))
    dst_rows = dst_tbl_el.findall(qn("w:tr"))

    for ri in range(min(len(src_rows), len(dst_rows))):
        src_tr = src_rows[ri]
        dst_tr = dst_rows[ri]

        # trPr
        src_trPr = src_tr.find(qn("w:trPr"))
        if src_trPr is not None:
            existing = dst_tr.find(qn("w:trPr"))
            if existing is not None:
                dst_tr.remove(existing)
            dst_tr.insert(0, deepcopy(src_trPr))

        src_cells = src_tr.findall(qn("w:tc"))
        dst_cells = dst_tr.findall(qn("w:tc"))
        for ci in range(min(len(src_cells), len(dst_cells))):
            src_tc = src_cells[ci]
            dst_tc = dst_cells[ci]

            # tcPr（含列宽、边框、底色、合并标记）
            src_tcPr = src_tc.find(qn("w:tcPr"))
            if src_tcPr is not None:
                existing = dst_tc.find(qn("w:tcPr"))
                if existing is not None:
                    dst_tc.remove(existing)
                dst_tc.insert(0, deepcopy(src_tcPr))

            # 复制段落格式（pPr + run rPr），保留文字
            src_paras = src_tc.findall(qn("w:p"))
            dst_paras = dst_tc.findall(qn("w:p"))
            for pi in range(min(len(src_paras), len(dst_paras))):
                src_pPr = src_paras[pi].find(qn("w:pPr"))
                dst_p   = dst_paras[pi]
                if src_pPr is not None:
                    ex_pPr = dst_p.find(qn("w:pPr"))
                    if ex_pPr is not None:
                        dst_p.remove(ex_pPr)
                    dst_p.insert(0, deepcopy(src_pPr))

                # 复制第一个 run 的 rPr 到所有 dst run
                src_runs = src_paras[pi].findall(qn("w:r"))
                if src_runs:
                    src_rPr = src_runs[0].find(qn("w:rPr"))
                    if src_rPr is not None:
                        for dst_r in dst_p.findall(qn("w:r")):
                            ex = dst_r.find(qn("w:rPr"))
                            if ex is not None:
                                dst_r.remove(ex)
                            dst_r.insert(0, deepcopy(src_rPr))


# =========================================================================== #
# copy_table 内部辅助                                                           #
# =========================================================================== #

def _apply_transform(tbl_el, src_table, transform: dict):
    """对 deepcopy 后的 w:tbl 元素应用 transform 变换。"""
    keep_cols = transform.get("columns")
    rename    = transform.get("rename_headers", {})
    max_rows  = transform.get("max_rows")
    filter_fn = transform.get("filter_rows")   # callable(row_dict) -> bool

    rows = tbl_el.findall(qn("w:tr"))
    if not rows:
        return tbl_el

    # ── 列过滤 ────────────────────────────────────────────────────────────── #
    if keep_cols is not None:
        keep_set = set(keep_cols)
        # tblGrid
        grid = tbl_el.find(qn("w:tblGrid"))
        if grid is not None:
            for i, gc in enumerate(list(grid.findall(qn("w:gridCol")))):
                if i not in keep_set:
                    grid.remove(gc)
        # 每行
        for tr in rows:
            cells = tr.findall(qn("w:tc"))
            for i, tc in enumerate(cells):
                if i not in keep_set:
                    tr.remove(tc)

    # ── 重命名列头（在第一行操作）────────────────────────────────────────── #
    if rename and rows:
        header_row = rows[0]
        for tc in header_row.findall(qn("w:tc")):
            for p in tc.findall(qn("w:p")):
                text = "".join(t.text or "" for t in p.iter(qn("w:t")))
                if text in rename:
                    # 清空现有 run，重新写入
                    for r in list(p.findall(qn("w:r"))):
                        p.remove(r)
                    new_r = OxmlElement("w:r")
                    new_t = OxmlElement("w:t")
                    new_t.text = rename[text]
                    new_r.append(new_t)
                    p.append(new_r)

    # ── 行过滤 + max_rows ─────────────────────────────────────────────────── #
    if (filter_fn is not None or max_rows is not None) and len(rows) > 1:
        data_rows = rows[1:]   # 跳过标题行

        # 重建原表格列头（用于 filter_rows 接收 dict）
        if filter_fn is not None:
            header_cells = rows[0].findall(qn("w:tc"))
            headers = [
                "".join(t.text or "" for t in tc.iter(qn("w:t")))
                for tc in header_cells
            ]

        kept = 0
        for tr in data_rows:
            if max_rows is not None and kept >= max_rows:
                tbl_el.remove(tr)
                continue
            if filter_fn is not None:
                cells = tr.findall(qn("w:tc"))
                row_dict = {
                    headers[i]: "".join(t.text or "" for t in c.iter(qn("w:t")))
                    for i, c in enumerate(cells) if i < len(headers)
                }
                if not filter_fn(row_dict):
                    tbl_el.remove(tr)
                    continue
            kept += 1

    return tbl_el


def _insert_table_at(dst_doc, new_tbl_el, dst_position: str,
                      author: str, as_tc_ins: bool,
                      next_available_tc_id: int | None = None) -> None:
    """将 new_tbl_el 插入 dst_doc 的指定位置。

    Args:
        next_available_tc_id: 若不为 None，replace_table 分支的 DEL TC ID
                              将从此值开始分配，避免与 INS 阶段的 ID 冲突。
    """
    body = dst_doc.element.body

    if dst_position.startswith("after_para:"):
        idx = int(dst_position.split(":")[1])
        para_el = dst_doc.paragraphs[idx]._element
        # 找到 para_el 在 body 中的位置，插在其后
        body_children = list(body)
        try:
            pos = body_children.index(para_el)
            body.insert(pos + 1, new_tbl_el)
        except ValueError:
            # para_el 不是 body 直接子元素（在表格内），降级到末尾
            body.append(new_tbl_el)

    elif dst_position.startswith("replace_table:"):
        idx = int(dst_position.split(":")[1])
        old_tbl_el = dst_doc.tables[idx]._tbl

        # TC DEL 旧表格（标记每行为删除）
        # 优先使用传入的 next_available_tc_id，避免与 INS 阶段 ID 重叠
        tc_id = next_available_tc_id if next_available_tc_id is not None else next_tc_id(dst_doc)
        date  = _utc_now()
        for tr in old_tbl_el.findall(qn("w:tr")):
            mark_row_as_deleted(tr, tc_id, author, date)
            tc_id += 1

        # 在旧表格后插入新表格
        body_children = list(body)
        try:
            pos = body_children.index(old_tbl_el)
            body.insert(pos + 1, new_tbl_el)
        except ValueError:
            body.append(new_tbl_el)

    else:
        raise ValueError(f"未知 dst_position 格式: {dst_position!r}"
                         f"（支持 'after_para:N' 或 'replace_table:N'）")
