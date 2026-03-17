"""
jt_note.py — 律所批注（Note）智能注入

适用于任何律所的任意批注格式，默认为竞天公诚的 [JT Note: ...]。
通过 cfg=DocConfig(...) 或直接参数可定制：
  - 批注开头/结尾标记（note_prefix / note_suffix）
  - 作者（author）
  - 高亮颜色（highlight）

三种注入模式：
  A. append_to_paragraph    — 在已有段落末尾追加 Note
  B. insert_paragraph       — 插入独立 Note 段落
  C. create_mixed_paragraph — 普通文字 + Note 混排的新段落
"""
from __future__ import annotations

from copy import deepcopy

from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from .constants import JT_AUTHOR, HIGHLIGHT_YELLOW, JT_NOTE_PREFIX, JT_NOTE_SUFFIX
from .tc_utils import _utc_now, next_tc_id, make_tc_tag, make_run


# --------------------------------------------------------------------------- #
# 参数解析辅助                                                                  #
# --------------------------------------------------------------------------- #

def _resolve(value, cfg_attr: str, fallback):
    """
    参数解析优先级：显式传参 > DocConfig 属性 > fallback 常量。
    value=None 表示"未传"，取 cfg 或 fallback。
    """
    if value is not None:
        return value
    from .config import DocConfig
    if isinstance(cfg_attr, str):  # cfg_attr is attribute name string
        pass  # handled below
    return fallback


def _get_cfg_attr(cfg, attr: str, fallback):
    if cfg is None:
        return fallback
    return getattr(cfg, attr, fallback)


# --------------------------------------------------------------------------- #
# 内部辅助                                                                      #
# --------------------------------------------------------------------------- #

def _wrap_note_text(note_text: str, prefix: str, suffix: str) -> str:
    """若未包裹 prefix…suffix，自动包裹。"""
    text = note_text.strip()
    if text.startswith(prefix.strip()):
        return text
    return f"{prefix}{text}{suffix}"


def _make_note_run(note_text: str, prefix: str, suffix: str,
                   highlight: str) -> OxmlElement:
    """构造带格式（B+I+HL）的 Note run。"""
    return make_run(
        _wrap_note_text(note_text, prefix, suffix),
        bold=True,
        italic=True,
        highlight=highlight,
    )


def _copy_pPr(source_para) -> OxmlElement | None:
    pPr = source_para._element.find(qn("w:pPr"))
    return deepcopy(pPr) if pPr is not None else None


# --------------------------------------------------------------------------- #
# 方式 A：已有段落末尾追加 Note                                                  #
# --------------------------------------------------------------------------- #

def append_to_paragraph(
    doc,
    paragraph_index: int,
    note_text: str,
    author: str | None = None,
    note_prefix: str | None = None,
    note_suffix: str | None = None,
    highlight: str | None = None,
    cfg=None,                           # DocConfig，可选
) -> None:
    """
    在 paragraph_index 段落末尾追加 Note，以 TC INS 形式注入。

    Args:
        doc:               python-docx Document
        paragraph_index:   目标段落索引
        note_text:         Note 正文（不含包裹标记，自动处理）
        author:            TC 作者（覆盖 cfg.author）
        note_prefix:       Note 开头标记，如 "[JT Note: "（覆盖 cfg.note_prefix）
        note_suffix:       Note 结尾标记，如 "]"
        highlight:         Word highlight 颜色，如 "yellow" / "cyan" / "green"
        cfg:               DocConfig 实例（提供批量默认值）
    """
    author     = author     or _get_cfg_attr(cfg, "author", JT_AUTHOR)
    note_prefix = note_prefix or _get_cfg_attr(cfg, "note_prefix", JT_NOTE_PREFIX)
    note_suffix = note_suffix or _get_cfg_attr(cfg, "note_suffix", JT_NOTE_SUFFIX)
    highlight  = highlight  or _get_cfg_attr(cfg, "note_highlight", HIGHLIGHT_YELLOW)

    para = doc.paragraphs[paragraph_index]
    tc_id = next_tc_id(doc)
    date  = _utc_now()

    prefix_space = " " if para.text and not para.text.endswith(" ") else ""

    ins = make_tc_tag("w:ins", tc_id, author, date)
    r   = make_run(
        prefix_space + _wrap_note_text(note_text, note_prefix, note_suffix),
        bold=True, italic=True, highlight=highlight,
    )
    ins.append(r)
    para._element.append(ins)


# --------------------------------------------------------------------------- #
# 方式 B：插入独立 Note 段落                                                    #
# --------------------------------------------------------------------------- #

def insert_paragraph(
    doc,
    after_index: int,
    note_text: str,
    author: str | None = None,
    inherit_style_from: int | None = None,
    note_prefix: str | None = None,
    note_suffix: str | None = None,
    highlight: str | None = None,
    cfg=None,
) -> None:
    """
    在 after_index 之后插入独立 Note 段落（整段为 TC INS）。

    Args:
        inherit_style_from: 从该段落索引复制 pPr（缩进/样式）
    """
    author      = author      or _get_cfg_attr(cfg, "author", JT_AUTHOR)
    note_prefix = note_prefix or _get_cfg_attr(cfg, "note_prefix", JT_NOTE_PREFIX)
    note_suffix = note_suffix or _get_cfg_attr(cfg, "note_suffix", JT_NOTE_SUFFIX)
    highlight   = highlight   or _get_cfg_attr(cfg, "note_highlight", HIGHLIGHT_YELLOW)

    paras    = doc.paragraphs
    ref_para = paras[after_index]
    tc_id    = next_tc_id(doc)
    date     = _utc_now()

    new_p = OxmlElement("w:p")
    if inherit_style_from is not None:
        pPr = _copy_pPr(paras[inherit_style_from])
        if pPr is not None:
            new_p.append(pPr)

    ins = make_tc_tag("w:ins", tc_id, author, date)
    ins.append(_make_note_run(note_text, note_prefix, note_suffix, highlight))
    new_p.append(ins)
    ref_para._element.addnext(new_p)


# --------------------------------------------------------------------------- #
# 方式 C：混合段落（普通文字 + Note 混排）                                      #
# --------------------------------------------------------------------------- #

def create_mixed_paragraph(
    doc,
    after_index: int,
    segments: list[tuple[str, bool]],
    author: str | None = None,
    style: str = "Normal",
    inherit_format_from: int | None = None,
    note_prefix: str | None = None,
    note_suffix: str | None = None,
    highlight: str | None = None,
    cfg=None,
) -> None:
    """
    在 after_index 之后插入混合段落。整段为一个 w:ins，
    普通文字 run 无格式，Note run 有 B+I+HL。

    Args:
        segments:  [(text, is_note), ...]
                   is_note=True → 自动加包裹标记 + B+I+HL
                   is_note=False → 普通文字，无格式修饰
        style:     段落样式名（当 inherit_format_from 为 None 时使用）
    """
    author      = author      or _get_cfg_attr(cfg, "author", JT_AUTHOR)
    note_prefix = note_prefix or _get_cfg_attr(cfg, "note_prefix", JT_NOTE_PREFIX)
    note_suffix = note_suffix or _get_cfg_attr(cfg, "note_suffix", JT_NOTE_SUFFIX)
    highlight   = highlight   or _get_cfg_attr(cfg, "note_highlight", HIGHLIGHT_YELLOW)

    paras    = doc.paragraphs
    ref_para = paras[after_index]
    tc_id    = next_tc_id(doc)
    date     = _utc_now()

    new_p = OxmlElement("w:p")

    if inherit_format_from is not None:
        pPr = _copy_pPr(paras[inherit_format_from])
        if pPr is not None:
            new_p.append(pPr)
    elif style:
        pPr    = OxmlElement("w:pPr")
        pStyle = OxmlElement("w:pStyle")
        pStyle.set(qn("w:val"), style)
        pPr.append(pStyle)
        new_p.append(pPr)

    ins = make_tc_tag("w:ins", tc_id, author, date)
    for text, is_note in segments:
        if not text:
            continue
        if is_note:
            ins.append(_make_note_run(text, note_prefix, note_suffix, highlight))
        else:
            ins.append(make_run(text))
    new_p.append(ins)
    ref_para._element.addnext(new_p)


# --------------------------------------------------------------------------- #
# 补救工具：修复已有 Note 格式                                                  #
# --------------------------------------------------------------------------- #

def fix_note_format(
    doc,
    para_range: tuple[int, int] | None = None,
    note_prefix: str | None = None,
    note_suffix: str | None = None,
    highlight: str | None = None,
    cfg=None,
) -> list[tuple[int, str]]:
    """
    扫描文档中所有含 Note 标记的 run，确保有 B+I+HL 格式。
    用于修复手动录入或粘贴导致格式丢失的情况。

    可通过 note_prefix 指定任意律所的 Note 标记前缀。

    Returns:
        修复的 (段落索引, run前20字) 列表
    """
    from .tc_utils import make_rPr

    note_prefix = note_prefix or _get_cfg_attr(cfg, "note_prefix", JT_NOTE_PREFIX)
    highlight   = highlight   or _get_cfg_attr(cfg, "note_highlight", HIGHLIGHT_YELLOW)
    # 检测关键字：去掉开头方括号，如 "[JT Note: " → "JT Note:"
    detect_str  = note_prefix.lstrip("[（(").rstrip(": ").strip()

    paras = doc.paragraphs
    start, end = para_range if para_range else (0, len(paras))
    fixed = []

    for i in range(start, min(end, len(paras))):
        para = paras[i]
        for run_el in para._element.iter(qn("w:r")):
            text = "".join(t.text or "" for t in run_el.findall(qn("w:t")))
            if detect_str not in text:
                continue

            rPr  = run_el.find(qn("w:rPr"))
            has_b  = rPr is not None and rPr.find(qn("w:b")) is not None
            has_i  = rPr is not None and rPr.find(qn("w:i")) is not None
            hl_el  = rPr.find(qn("w:highlight")) if rPr is not None else None
            has_hl = hl_el is not None and hl_el.get(qn("w:val")) == highlight

            if has_b and has_i and has_hl:
                continue

            new_rPr = make_rPr(bold=True, italic=True, highlight=highlight,
                                base_rPr=rPr)
            if rPr is not None:
                run_el.remove(rPr)
            run_el.insert(0, new_rPr)
            fixed.append((i, text[:20]))

    return fixed


# 向后兼容别名
fix_jt_note_format = fix_note_format
