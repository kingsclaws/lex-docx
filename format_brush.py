"""
format_brush.py — 格式刷

解决 TC INS 注入后新段落丢失缩进、段间距、样式的问题。
从参考段落复制 w:pPr 子元素到目标段落，支持选择性复制。
"""
from __future__ import annotations

from copy import deepcopy
from typing import Sequence

from docx.oxml.ns import qn

# 支持复制的 pPr 子元素标签
_COPY_MAP = {
    "indent":   qn("w:ind"),        # 缩进
    "spacing":  qn("w:spacing"),    # 段间距
    "style":    qn("w:pStyle"),     # 段落样式
    "numPr":    qn("w:numPr"),      # 编号/列表
    "jc":       qn("w:jc"),         # 对齐方式
    "outlineLvl": qn("w:outlineLvl"), # 大纲级别
}


# --------------------------------------------------------------------------- #
# 主接口                                                                        #
# --------------------------------------------------------------------------- #

def apply(
    doc,
    target_indices: Sequence[int],
    reference_index: int,
    copy: list[str] | None = None,
) -> list[int]:
    """
    从 reference_index 段落复制格式到 target_indices 中的各段落。

    Args:
        doc:              python-docx Document 对象
        target_indices:   需要修复格式的段落索引列表
        reference_index:  格式正确的参考段落索引
        copy:             选择性复制，支持 "indent" / "spacing" / "style" /
                          "numPr" / "jc" / "outlineLvl"
                          默认 ["indent", "spacing", "style"]

    Returns:
        实际修改的段落索引列表
    """
    if copy is None:
        copy = ["indent", "spacing", "style"]

    paras = doc.paragraphs
    ref_para = paras[reference_index]
    ref_pPr = ref_para._element.find(qn("w:pPr"))

    if ref_pPr is None:
        return []   # 参考段落没有 pPr，无可复制

    modified = []
    for idx in target_indices:
        if idx < 0 or idx >= len(paras):
            continue
        tgt_para = paras[idx]
        tgt_pPr = tgt_para._element.get_or_add_pPr()

        for attr in copy:
            tag = _COPY_MAP.get(attr)
            if tag is None:
                continue
            src_child = ref_pPr.find(tag)
            if src_child is None:
                continue
            # 删除目标中已有的同名元素，再插入复制版本
            existing = tgt_pPr.find(tag)
            if existing is not None:
                tgt_pPr.remove(existing)
            tgt_pPr.append(deepcopy(src_child))

        modified.append(idx)

    return modified


def auto_fix(
    doc,
    para_range: tuple[int, int] | None = None,
    template_doc=None,
    copy: list[str] | None = None,
) -> list[int]:
    """
    按 style name 自动匹配参考段落，批量修复缩进/间距。

    逻辑：
    1. 扫描 template_doc（或文档本身 para_range 之外的段落）建立
       style_name → pPr 的参考映射（取每种样式的第一个有 pPr 的段落）
    2. 对 para_range 内的每个段落，用其 style name 查参考，
       将 indent / spacing 对齐到参考值

    Args:
        doc:          python-docx Document 对象
        para_range:   需要修复的段落范围 (start, end)，默认处理全文
        template_doc: 可选，从另一个文档提取参考样式
        copy:         默认 ["indent", "spacing"]（auto_fix 不改 style 本身）

    Returns:
        实际修改的段落索引列表
    """
    if copy is None:
        copy = ["indent", "spacing"]

    paras = doc.paragraphs
    start, end = para_range if para_range else (0, len(paras))

    # 建立 style → 参考 pPr 映射
    style_refs: dict[str, object] = {}
    source_paras = template_doc.paragraphs if template_doc else paras

    for para in source_paras:
        style_name = para.style.name if para.style else "Normal"
        if style_name not in style_refs:
            pPr = para._element.find(qn("w:pPr"))
            if pPr is not None:
                # 确保至少有 ind 或 spacing 可用
                has_useful = any(
                    pPr.find(qn(f"w:{t}")) is not None
                    for t in ("ind", "spacing")
                )
                if has_useful:
                    style_refs[style_name] = pPr

    modified = []
    for i in range(start, min(end, len(paras))):
        para = paras[i]
        style_name = para.style.name if para.style else "Normal"
        ref_pPr = style_refs.get(style_name)
        if ref_pPr is None:
            continue

        tgt_pPr = para._element.get_or_add_pPr()
        changed = False
        for attr in copy:
            tag = _COPY_MAP.get(attr)
            if tag is None:
                continue
            src_child = ref_pPr.find(tag)
            if src_child is None:
                continue
            existing = tgt_pPr.find(tag)
            # 比较 XML 文本，只有不一致才修改
            src_xml = _el_to_str(src_child)
            if existing is not None and _el_to_str(existing) == src_xml:
                continue
            if existing is not None:
                tgt_pPr.remove(existing)
            tgt_pPr.append(deepcopy(src_child))
            changed = True

        if changed:
            modified.append(i)

    return modified


# --------------------------------------------------------------------------- #
# 辅助                                                                          #
# --------------------------------------------------------------------------- #

def get_pPr_summary(doc, para_range: tuple[int, int] | None = None) -> list[dict]:
    """
    调试用：返回段落范围内每段的 style / ind / spacing 摘要。
    方便确认哪些段落格式不一致。
    """
    paras = doc.paragraphs
    start, end = para_range if para_range else (0, len(paras))
    result = []
    for i in range(start, min(end, len(paras))):
        para = paras[i]
        pPr = para._element.find(qn("w:pPr"))
        ind = pPr.find(qn("w:ind")) if pPr is not None else None
        spacing = pPr.find(qn("w:spacing")) if pPr is not None else None
        result.append({
            "index": i,
            "style": para.style.name if para.style else "Normal",
            "text_preview": para.text[:40],
            "ind_left": ind.get(qn("w:left")) if ind is not None else None,
            "ind_hanging": ind.get(qn("w:hanging")) if ind is not None else None,
            "spacing_before": spacing.get(qn("w:before")) if spacing is not None else None,
            "spacing_after": spacing.get(qn("w:after")) if spacing is not None else None,
        })
    return result


def _el_to_str(el) -> str:
    """将 lxml element 序列化为字符串，用于比较是否相同。"""
    from lxml import etree
    return etree.tostring(el, encoding="unicode")


def set_outline_level(
    doc,
    target_indices: Sequence[int],
    level: int | None,
) -> list[int]:
    """
    直接设置段落的 w:outlineLvl 值。

    Args:
        doc:            python-docx Document 对象
        target_indices: 目标段落索引列表
        level:          1–9（对应 OOXML 0–8），或 None / 0 = 清除（变为正文级别）

    Returns:
        实际修改的段落索引列表

    说明：
        - w:outlineLvl 独立于 Heading 1/2/3 样式，控制 Word 导航窗格中的大纲层级
        - val=0 → 大纲1级，val=8 → 大纲9级，val=9 / 无此元素 → 正文（不出现在大纲）
        - 常见用途：自定义标题样式希望出现在导航窗格时，手动设置此值
    """
    tag = qn("w:outlineLvl")
    paras = doc.paragraphs

    # 转换用户级别（1-9）到 OOXML 值（0-8），None/0 表示清除
    if level is None or level <= 0:
        ooxml_val = None   # 清除模式
    else:
        ooxml_val = str(min(level - 1, 8))

    modified = []
    for idx in target_indices:
        if idx < 0 or idx >= len(paras):
            continue
        para = paras[idx]
        pPr = para._element.get_or_add_pPr()
        existing = pPr.find(tag)

        if ooxml_val is None:
            # 清除大纲级别
            if existing is not None:
                pPr.remove(existing)
                modified.append(idx)
        else:
            if existing is not None:
                if existing.get(qn("w:val")) == ooxml_val:
                    continue   # 已是目标值，跳过
                pPr.remove(existing)
            from lxml import etree
            el = etree.SubElement(pPr, tag)
            el.set(qn("w:val"), ooxml_val)
            modified.append(idx)

    return modified


def extract_style_rPr_map(doc) -> dict:
    """
    从文档 styles.xml 提取每个 style 的 w:rPr，返回 {style_name: rPr_element}。

    返回值可直接传给 tc_ins_text(style_rPr_map=...) 或存入 DocConfig.style_rPr_map。
    lxml element 版本（本函数返回的）可直接 deepcopy 使用。

    示例：
        style_map = format_brush.extract_style_rPr_map(doc)
        tc_utils.tc_ins_text(para, "新文字", tc_id, "JT",
                             inherit_rPr="auto", style_rPr_map=style_map)
    """
    result: dict = {}
    for style in doc.styles:
        rPr = style._element.find(qn("w:rPr"))
        if rPr is not None:
            result[style.name] = rPr   # 不 deepcopy，调用方自行 deepcopy
    return result
