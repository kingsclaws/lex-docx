"""
inject_engine.py — 批量注入引擎（FR-7）

将结构化注入计划（InjectPlan）转换为一系列 DOCX 操作，
替代手写的 mapping dict 脚本。适合每个主体章节的批量替换场景。

典型用法：
    from lex_docx import inject_engine, DocConfig

    cfg = DocConfig(author="JT")
    plan = inject_engine.InjectPlan(
        doc_path="report.docx",
        out_path="report_reviewed.docx",
        tables=[
            inject_engine.TableFill(
                table_index=8,
                data={"企业名称": "临港资管", "法定代表人": "张三"},
                mode="kv",
            ),
            inject_engine.TableFill(
                table_index=12,
                data=[{"股东": "临港集团", "出资额": "1060万", "比例": "100%"}],
                mode="rows",
                auto_adjust=True,
            ),
        ],
        jt_notes={
            180: "待获取完整工商登记资料后确认。",
            "治理结构": "待获取公司章程后核实。",
        },
        auto_cleanup=True,
        run_lint=True,
    )
    result = inject_engine.execute(plan, cfg)
    print(result.summary())
"""
from __future__ import annotations

from dataclasses import dataclass, field
from pathlib import Path
from typing import Any


# =========================================================================== #
# 数据结构                                                                      #
# =========================================================================== #

@dataclass
class TableFill:
    """
    单张表格的注入描述。

    mode:
      "kv"   — KV 表（基本信息类），data 为 dict
      "rows" — 行数据表（股东表/诉讼表等），data 为 List[dict]

    column_mapping: {源列名: 目标列名} 或 {源列名: "auto_number"}
    key_columns:    四列布局时的 key 列索引列表，如 [0, 2]
    auto_adjust:    行数不足时自动 TC INS 增行，多余行自动 TC DEL
    header_row_index: 标题行位置（默认 0）
    """
    table_index: int
    data: dict | list[dict]
    mode: str = "kv"                       # "kv" | "rows"
    column_mapping: dict | None = None
    key_columns: list[int] | None = None
    auto_adjust: bool = True
    header_row_index: int = 0


@dataclass
class InjectResult:
    tables: list[dict] = field(default_factory=list)      # [{table_index, filled, mode}, ...]
    notes: list[dict] = field(default_factory=list)        # [{key, status}, ...]
    cleanup: dict = field(default_factory=dict)            # {"empty": [...], "orphan_numbering": [...]}
    lint: list | None = None                               # List[LintResult] or None

    def summary(self) -> str:
        lines = []
        total_filled = sum(t.get("filled", 0) for t in self.tables)
        lines.append(f"表格: {len(self.tables)} 张, 共填充 {total_filled} 行/格")
        lines.append(f"JT Note: 注入 {len(self.notes)} 条")
        if self.cleanup:
            empty = len(self.cleanup.get("empty", []))
            orphan = len(self.cleanup.get("orphan_numbering", []))
            lines.append(f"清理: 空段落 {empty} 处, 孤儿编号 {orphan} 处")
        if self.lint is not None:
            passed = sum(1 for r in self.lint if r.passed)
            lines.append(f"Lint: {passed}/{len(self.lint)} 条规则通过")
            for r in self.lint:
                if not r.passed:
                    lines.append(f"  ❌ {r.rule}: {r.detail}")
        return "\n".join(lines)


@dataclass
class InjectPlan:
    """
    注入计划。

    doc_path:      源 DOCX 文件路径
    out_path:      输出路径（None = 覆盖 doc_path）
    target_range:  段落范围（start, end），用于 cleanup 和 lint 的 check_range
    tables:        TableFill 列表
    jt_notes:      {para_index_or_keyword: note_text}
                   int key → 在该段落末尾 append
                   str key → 在第一个包含该文字的段落末尾 append
    auto_cleanup:  注入后自动运行 cleanup.cleanup_all()
    run_lint:      注入后自动运行 lint.check()
    """
    doc_path: str
    out_path: str | None = None
    target_range: tuple[int, int] | None = None
    tables: list[TableFill] = field(default_factory=list)
    jt_notes: dict[int | str, str] = field(default_factory=dict)
    auto_cleanup: bool = True
    run_lint: bool = True


# =========================================================================== #
# 执行引擎                                                                      #
# =========================================================================== #

def execute(plan: InjectPlan, cfg=None) -> InjectResult:
    """
    执行注入计划，返回 InjectResult。

    Args:
        plan: InjectPlan 实例
        cfg:  DocConfig（提供 author / note_prefix 等）
    """
    from docx import Document
    from . import table_ops, jt_note, cleanup, lint

    result = InjectResult()
    doc = Document(plan.doc_path)

    # ── 1. 表格注入 ────────────────────────────────────────────────────────── #
    for tf in plan.tables:
        try:
            if tf.mode == "kv":
                if not isinstance(tf.data, dict):
                    raise ValueError(f"mode='kv' 要求 data 为 dict，实际为 {type(tf.data)}")
                filled = table_ops.fill_kv_table(
                    doc, tf.table_index, tf.data,
                    key_columns=tf.key_columns,
                    cfg=cfg,
                )
            else:  # "rows"
                if not isinstance(tf.data, list):
                    raise ValueError(f"mode='rows' 要求 data 为 list，实际为 {type(tf.data)}")
                if tf.auto_adjust:
                    table_ops.adjust_rows(
                        doc, tf.table_index, len(tf.data),
                        header_row_index=tf.header_row_index,
                        cfg=cfg,
                    )
                filled = table_ops.fill_table(
                    doc, tf.table_index, tf.data,
                    column_mapping=tf.column_mapping,
                    header_row_index=tf.header_row_index,
                    cfg=cfg,
                )
            result.tables.append({
                "table_index": tf.table_index,
                "mode": tf.mode,
                "filled": filled,
                "status": "ok",
            })
        except Exception as e:
            result.tables.append({
                "table_index": tf.table_index,
                "mode": tf.mode,
                "filled": 0,
                "status": f"error: {e}",
            })

    # ── 2. JT Note 注入 ────────────────────────────────────────────────────── #
    paras = doc.paragraphs
    for key, note_text in plan.jt_notes.items():
        try:
            if isinstance(key, int):
                jt_note.append_to_paragraph(doc, key, note_text, cfg=cfg)
                result.notes.append({"key": key, "status": "ok"})
            else:
                # str key → 找第一个包含该文字的段落
                found = False
                for i, para in enumerate(paras):
                    if key in para.text:
                        jt_note.append_to_paragraph(doc, i, note_text, cfg=cfg)
                        result.notes.append({"key": key, "para_index": i, "status": "ok"})
                        found = True
                        break
                if not found:
                    result.notes.append({"key": key, "status": f"未找到包含 {key!r} 的段落"})
        except Exception as e:
            result.notes.append({"key": key, "status": f"error: {e}"})

    # ── 3. 保存 ────────────────────────────────────────────────────────────── #
    out_path = plan.out_path or plan.doc_path
    doc.save(out_path)

    # ── 4. 清理（在保存后重新加载，避免索引偏移影响后续操作）──────────────────── #
    if plan.auto_cleanup:
        doc2 = Document(out_path)
        cr = cleanup.cleanup_all(
            doc2,
            para_range=plan.target_range,
            cfg=cfg,
        )
        result.cleanup = cr
        doc2.save(out_path)

    # ── 5. Lint ────────────────────────────────────────────────────────────── #
    if plan.run_lint:
        lint_cfg = cfg
        if cfg is not None and plan.target_range is not None:
            # 将 check_range 注入 lint config
            if hasattr(cfg, "to_lint_config"):
                lint_cfg = cfg.to_lint_config(extra={"check_range": plan.target_range})
        result.lint = lint.check(out_path, config=lint_cfg)

    return result
