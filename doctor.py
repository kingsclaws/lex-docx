"""
doctor.py — 格式诊断与自动修复

Rules:
  D01  font_missing              — run 缺少 rFonts.eastAsia（会回退到文档默认字体）
  D02  font_mismatch             — run 的 rFonts.eastAsia 与标准字体不符
  D03  double_numbering          — 段落有 numPr 且文本以手动编号开头（不自动修复）
  D04  outline_leak              — 无显式 outlineLvl 但样式继承链有 outlineLvl 0-2
  D05  sibling_numpr_gap         — 同 styleId 段落中部分缺少 numPr
  D06  invalid_style_id          — pStyle 引用不存在的 styleId
  D07  toc_u_switch              — TOC 字段含 \\u 开关（绕过显式 outlineLvl）
  D08  heading_font_inconsistent — 同 styleId 标题字体/字号不一致

auto_fix 支持：D01 D02 D04 D05 D07 D08
不自动修复：D03（需人工判断） D06（需人工确认）
"""
from __future__ import annotations

import re
from collections import Counter, defaultdict
from copy import deepcopy
from dataclasses import dataclass, field
from typing import Any

from docx.oxml.ns import qn
from lxml import etree


# --------------------------------------------------------------------------- #
# 数据结构                                                                      #
# --------------------------------------------------------------------------- #

@dataclass
class Standards:
    """由调用方（agent/CLI）传入的格式标准，不在代码中硬编码默认值。"""
    font: str | None = None           # 标准 rFonts.eastAsia（None = 从文档推断）
    ascii_font: str | None = None     # 标准 rFonts.ascii / hAnsi（None = 同 font）
    font_size: float | None = None    # 标准字号 pt（None = 不检查字号）
    toc_levels: tuple[int, int] = (1, 3)  # TOC 应收录的大纲级别范围
    double_num_patterns: list[str] | None = None  # 自定义手动编号正则（None = 用内置）
    footer_blacklist: list[str] | None = None     # footer 关键词黑名单（如 ["Auspicious"]）


@dataclass
class Issue:
    rule: str                   # "D01" ~ "D08"
    severity: str               # "error" | "warning"
    para: int                   # 段落索引
    run: int | None             # run 索引（None = 段落级）
    detail: str
    auto_fix: bool
    extra: dict = field(default_factory=dict)   # 修复时用的附加信息


@dataclass
class CheckResult:
    issues: list[Issue]
    inferred_font: str | None   # 从文档推断的标准字体（供 fix 阶段使用）

    def summary(self) -> dict:
        from collections import Counter
        by_rule = Counter(i.rule for i in self.issues)
        errors = sum(1 for i in self.issues if i.severity == "error")
        warnings = sum(1 for i in self.issues if i.severity == "warning")
        return {
            "total": len(self.issues),
            "errors": errors,
            "warnings": warnings,
            "by_rule": dict(by_rule),
        }

    def to_dict(self) -> dict:
        return {
            "summary": self.summary(),
            "inferred_font": self.inferred_font,
            "issues": [
                {
                    "rule": i.rule,
                    "severity": i.severity,
                    "para": i.para,
                    "run": i.run,
                    "detail": i.detail,
                    "auto_fix": i.auto_fix,
                }
                for i in self.issues
            ],
        }


@dataclass
class FixResult:
    fixed: int
    skipped: int
    dry_run: bool
    log: list[str]


# --------------------------------------------------------------------------- #
# 样式信息构建（styleId 索引，含 basedOn 继承链）                               #
# --------------------------------------------------------------------------- #

def _build_style_info(doc) -> dict:
    """
    返回两个 dict：
        by_name[style_name]  → StyleInfo
        by_id[style_id]      → StyleInfo

    StyleInfo 结构：
    {
        "id": str,
        "name": str,
        "based_on_name": str | None,
        "own_outline_lvl": int | None,     # 本层 pPr.outlineLvl
        "effective_outline_lvl": int | None,  # 含继承链
        "own_numpr": dict | None,          # {numId, ilvl}
        "effective_numpr": dict | None,
        "own_rpr": {"font_ea": str|None, "font_ascii": str|None, "font_size": float|None},
    }
    """
    raw: dict[str, dict] = {}         # name → raw (不含继承)
    by_id: dict[str, str] = {}        # styleId → name
    based_on_map: dict[str, str] = {} # name → parent_name

    for style in doc.styles:
        el = style.element
        sid = el.get(qn("w:styleId"), "")
        name = style.name or sid
        by_id[sid] = name

        parent_el = el.find(qn("w:basedOn"))
        parent_name = None
        if parent_el is not None:
            pid = parent_el.get(qn("w:val"), "")
            parent_name = by_id.get(pid, pid)
            based_on_map[name] = parent_name

        # pPr
        pPr = el.find(qn("w:pPr"))
        own_outline = None
        own_numpr = None
        if pPr is not None:
            ol = pPr.find(qn("w:outlineLvl"))
            if ol is not None:
                try:
                    own_outline = int(ol.get(qn("w:val"), 9))
                except (ValueError, TypeError):
                    pass
            np = pPr.find(qn("w:numPr"))
            if np is not None:
                ni = np.find(qn("w:numId"))
                il = np.find(qn("w:ilvl"))
                if ni is not None:
                    own_numpr = {
                        "numId": ni.get(qn("w:val")),
                        "ilvl": il.get(qn("w:val")) if il is not None else "0",
                    }

        # rPr
        rPr = el.find(qn("w:rPr"))
        own_rpr: dict = {"font_ea": None, "font_ascii": None, "font_size": None}
        if rPr is not None:
            rFonts = rPr.find(qn("w:rFonts"))
            if rFonts is not None:
                own_rpr["font_ea"] = rFonts.get(qn("w:eastAsia"))
                own_rpr["font_ascii"] = rFonts.get(qn("w:ascii"))
            sz = rPr.find(qn("w:sz"))
            if sz is not None:
                try:
                    own_rpr["font_size"] = int(sz.get(qn("w:val"), 0)) / 2
                except (ValueError, TypeError):
                    pass

        raw[name] = {
            "id": sid,
            "name": name,
            "based_on_name": parent_name,
            "own_outline_lvl": own_outline,
            "own_numpr": own_numpr,
            "own_rpr": own_rpr,
        }

    # 解析继承链
    def _resolve_outline(name: str, visited: set | None = None) -> int | None:
        if visited is None:
            visited = set()
        if name in visited:
            return None
        visited.add(name)
        info = raw.get(name, {})
        if info.get("own_outline_lvl") is not None:
            return info["own_outline_lvl"]
        parent = info.get("based_on_name")
        return _resolve_outline(parent, visited) if parent else None

    def _resolve_numpr(name: str, visited: set | None = None) -> dict | None:
        if visited is None:
            visited = set()
        if name in visited:
            return None
        visited.add(name)
        info = raw.get(name, {})
        if info.get("own_numpr") is not None:
            return info["own_numpr"]
        parent = info.get("based_on_name")
        return _resolve_numpr(parent, visited) if parent else None

    result_by_name: dict[str, dict] = {}
    for name, info in raw.items():
        resolved = dict(info)
        resolved["effective_outline_lvl"] = _resolve_outline(name)
        resolved["effective_numpr"] = _resolve_numpr(name)
        result_by_name[name] = resolved

    result_by_id = {info["id"]: info for info in result_by_name.values() if info["id"]}
    return {"by_name": result_by_name, "by_id": result_by_id}


def _para_style_name(para) -> str:
    return para.style.name if para.style else "Normal"


def _para_own_outline(para) -> int | None:
    pPr = para._element.find(qn("w:pPr"))
    if pPr is None:
        return None
    ol = pPr.find(qn("w:outlineLvl"))
    if ol is None:
        return None
    try:
        return int(ol.get(qn("w:val"), 9))
    except (ValueError, TypeError):
        return None


def _para_own_numpr(para) -> dict | None:
    pPr = para._element.find(qn("w:pPr"))
    if pPr is None:
        return None
    np = pPr.find(qn("w:numPr"))
    if np is None:
        return None
    ni = np.find(qn("w:numId"))
    il = np.find(qn("w:ilvl"))
    if ni is None:
        return None
    return {
        "numId": ni.get(qn("w:val")),
        "ilvl": il.get(qn("w:val")) if il is not None else "0",
    }


def _effective_numpr(para, style_info: dict) -> dict | None:
    own = _para_own_numpr(para)
    if own is not None:
        return own
    return style_info["by_name"].get(_para_style_name(para), {}).get("effective_numpr")


# --------------------------------------------------------------------------- #
# 字体推断                                                                     #
# --------------------------------------------------------------------------- #

def _infer_standard_font(doc) -> str | None:
    """统计文档中所有 run 的 rFonts.eastAsia，返回最频繁的值。"""
    counter: Counter = Counter()
    for para in doc.paragraphs:
        for run_el in para._element.iter(qn("w:r")):
            rPr = run_el.find(qn("w:rPr"))
            if rPr is None:
                continue
            rFonts = rPr.find(qn("w:rFonts"))
            if rFonts is None:
                continue
            ea = rFonts.get(qn("w:eastAsia"))
            if ea:
                counter[ea] += 1
    return counter.most_common(1)[0][0] if counter else None


# --------------------------------------------------------------------------- #
# 规则检查函数                                                                  #
# --------------------------------------------------------------------------- #

def _iter_para_runs(para):
    """遍历段落内所有 w:r（含 w:ins 内的 TC INS run），返回 (run_idx, run_el)。"""
    idx = 0
    for el in para._element.iter(qn("w:r")):
        yield idx, el
        idx += 1


def _check_d01_d02(doc, standards: Standards, para_range: tuple | None) -> list[Issue]:
    """D01: 缺 eastAsia；D02: eastAsia 与标准不符。"""
    issues = []
    paras = doc.paragraphs
    start, end = para_range or (0, len(paras))
    std_font = standards.font

    for pi in range(start, min(end, len(paras))):
        para = paras[pi]
        for ri, run_el in _iter_para_runs(para):
            rPr = run_el.find(qn("w:rPr"))
            # D01: 无 rPr 或无 rFonts 或无 eastAsia
            if rPr is None:
                issues.append(Issue(
                    rule="D01", severity="error", para=pi, run=ri,
                    detail="run 无 rPr，缺少 rFonts.eastAsia",
                    auto_fix=True,
                ))
                continue
            rFonts = rPr.find(qn("w:rFonts"))
            if rFonts is None:
                issues.append(Issue(
                    rule="D01", severity="error", para=pi, run=ri,
                    detail="run 缺少 rFonts 元素，eastAsia 未定义",
                    auto_fix=True,
                ))
                continue
            ea = rFonts.get(qn("w:eastAsia"))
            if ea is None:
                issues.append(Issue(
                    rule="D01", severity="error", para=pi, run=ri,
                    detail="run 缺少 rFonts.eastAsia 属性",
                    auto_fix=True,
                ))
            elif std_font and ea != std_font:
                issues.append(Issue(
                    rule="D02", severity="warning", para=pi, run=ri,
                    detail=f"rFonts.eastAsia={ea!r}，期望 {std_font!r}",
                    auto_fix=True,
                    extra={"current_font": ea},
                ))

    return issues


_DOUBLE_NUM_PATTERNS = [
    r"^[（(][一二三四五六七八九十百\d]+[）)]",  # （一）（1）
    r"^\d+[.、．]\s*\S",                         # 1. 1、
    r"^[一二三四五六七八九十]+[、．.]\s*\S",      # 一、二、
    r"^第[一二三四五六七八九十百\d]+[章节条款]",  # 第一章
]

def _check_d03(doc, style_info: dict, para_range: tuple | None,
               custom_patterns: list[str] | None) -> list[Issue]:
    """D03: 段落有 numPr 且文本以手动编号开头。"""
    patterns = [re.compile(p) for p in (custom_patterns or _DOUBLE_NUM_PATTERNS)]
    issues = []
    paras = doc.paragraphs
    start, end = para_range or (0, len(paras))

    for pi in range(start, min(end, len(paras))):
        para = paras[pi]
        eff_np = _effective_numpr(para, style_info)
        if eff_np is None:
            continue
        text = para.text.strip()
        if not text:
            continue
        for pat in patterns:
            if pat.match(text):
                issues.append(Issue(
                    rule="D03", severity="warning", para=pi, run=None,
                    detail=f"段落有 numPr（numId={eff_np['numId']}），文本同时含手动编号: {text[:40]!r}",
                    auto_fix=False,
                ))
                break

    return issues


def _check_d04(doc, style_info: dict, para_range: tuple | None) -> list[Issue]:
    """D04: 无显式 outlineLvl 但样式继承链有 outlineLvl 0-2（会被 TOC \\o 收录）。"""
    issues = []
    paras = doc.paragraphs
    start, end = para_range or (0, len(paras))
    by_name = style_info["by_name"]

    for pi in range(start, min(end, len(paras))):
        para = paras[pi]
        if _para_own_outline(para) is not None:
            continue   # 有显式设置，正常
        style_name = _para_style_name(para)
        eff = by_name.get(style_name, {}).get("effective_outline_lvl")
        if eff is not None and 0 <= eff <= 2:
            issues.append(Issue(
                rule="D04", severity="warning", para=pi, run=None,
                detail=f"样式 {style_name!r} 继承 outlineLvl={eff}，段落无显式覆盖，将被 TOC \\o 收录",
                auto_fix=True,
                extra={"inherited_outline": eff, "style": style_name},
            ))

    return issues


def _check_d05(doc, style_info: dict, para_range: tuple | None) -> list[Issue]:
    """D05: 同 styleId 段落中部分缺少 numPr（编号断裂）。"""
    paras = doc.paragraphs
    start, end = para_range or (0, len(paras))
    by_name = style_info["by_name"]

    # 按 styleId 分组
    groups: dict[str, list[int]] = defaultdict(list)
    for pi in range(start, min(end, len(paras))):
        para = paras[pi]
        sid = by_name.get(_para_style_name(para), {}).get("id", "")
        if sid:
            groups[sid].append(pi)

    issues = []
    for sid, indices in groups.items():
        if len(indices) < 2:
            continue
        has_np = [i for i in indices if _para_own_numpr(paras[i]) is not None]
        no_np  = [i for i in indices if _para_own_numpr(paras[i]) is None]
        if not has_np or not no_np:
            continue
        # 多数有 numPr，少数没有 → 报告没有的
        if len(has_np) >= len(no_np):
            # 找一个代表性的 numPr 用于修复
            ref_numpr = _para_own_numpr(paras[has_np[0]])
            for pi in no_np:
                issues.append(Issue(
                    rule="D05", severity="warning", para=pi, run=None,
                    detail=f"段落缺少 numPr（同 styleId {sid!r} 的其他 {len(has_np)} 个段落有 numPr）",
                    auto_fix=True,
                    extra={"ref_numpr": ref_numpr, "style_id": sid},
                ))

    return issues


def _check_d06(doc, style_info: dict, para_range: tuple | None) -> list[Issue]:
    """D06: pStyle 引用的 styleId 在 styles.xml 中不存在。"""
    issues = []
    paras = doc.paragraphs
    start, end = para_range or (0, len(paras))
    valid_ids = set(style_info["by_id"].keys())

    for pi in range(start, min(end, len(paras))):
        para = paras[pi]
        pPr = para._element.find(qn("w:pPr"))
        if pPr is None:
            continue
        pStyle = pPr.find(qn("w:pStyle"))
        if pStyle is None:
            continue
        sid = pStyle.get(qn("w:val"), "")
        if sid and sid not in valid_ids:
            # 尝试模糊匹配
            candidates = [k for k in valid_ids if sid.lower() in k.lower() or k.lower() in sid.lower()]
            issues.append(Issue(
                rule="D06", severity="error", para=pi, run=None,
                detail=f"pStyle 引用不存在的 styleId={sid!r}，可能候选: {candidates[:3]}",
                auto_fix=False,
                extra={"invalid_id": sid, "candidates": candidates[:3]},
            ))

    return issues


def _check_d07(doc) -> list[Issue]:
    """D07: TOC 字段含 \\u 开关。"""
    issues = []
    for pi, para in enumerate(doc.paragraphs):
        for run_el in para._element.iter(qn("w:r")):
            instr = run_el.find(qn("w:instrText"))
            if instr is None:
                continue
            text = instr.text or ""
            if text.strip().startswith("TOC") and r"\u" in text:
                issues.append(Issue(
                    rule="D07", severity="warning", para=pi, run=None,
                    detail=f"TOC 字段含 \\u 开关（将绕过显式 outlineLvl）: {text.strip()[:80]!r}",
                    auto_fix=True,
                    extra={"instr_text": text},
                ))

    return issues


def _check_d08(doc, standards: Standards, style_info: dict,
               para_range: tuple | None) -> list[Issue]:
    """D08: 同 styleId 标题段落字体/字号不一致。"""
    paras = doc.paragraphs
    start, end = para_range or (0, len(paras))
    by_name = style_info["by_name"]

    # 只检查有 outlineLvl（0-8）的样式的段落，即标题类
    heading_styles: set[str] = {
        name for name, info in by_name.items()
        if info.get("effective_outline_lvl") is not None
    }

    # 按 (styleId, effective_outline) 分组，收集 run 字体/字号
    groups: dict[str, list[tuple[int, str | None, float | None]]] = defaultdict(list)
    for pi in range(start, min(end, len(paras))):
        para = paras[pi]
        sname = _para_style_name(para)
        if sname not in heading_styles:
            continue
        sid = by_name.get(sname, {}).get("id", sname)
        for _, run_el in _iter_para_runs(para):
            rPr = run_el.find(qn("w:rPr"))
            ea = None
            sz = None
            if rPr is not None:
                rFonts = rPr.find(qn("w:rFonts"))
                if rFonts is not None:
                    ea = rFonts.get(qn("w:eastAsia"))
                sz_el = rPr.find(qn("w:sz"))
                if sz_el is not None:
                    try:
                        sz = int(sz_el.get(qn("w:val"), 0)) / 2
                    except (ValueError, TypeError):
                        pass
            groups[sid].append((pi, ea, sz))

    issues = []
    for sid, entries in groups.items():
        fonts = [e[1] for e in entries if e[1]]
        sizes = [e[2] for e in entries if e[2]]
        if len(set(fonts)) > 1:
            majority_font = Counter(fonts).most_common(1)[0][0]
            outliers = {e[0] for e in entries if e[1] and e[1] != majority_font}
            for pi in outliers:
                issues.append(Issue(
                    rule="D08", severity="warning", para=pi, run=None,
                    detail=f"标题样式 {sid!r} 字体不一致（多数为 {majority_font!r}）",
                    auto_fix=True,
                    extra={"majority_font": majority_font, "style_id": sid},
                ))
        if len(set(sizes)) > 1:
            majority_size = Counter(sizes).most_common(1)[0][0]
            outliers = {e[0] for e in entries if e[2] and abs(e[2] - majority_size) > 0.1}
            for pi in outliers:
                if not any(i.rule == "D08" and i.para == pi for i in issues):
                    issues.append(Issue(
                        rule="D08", severity="warning", para=pi, run=None,
                        detail=f"标题样式 {sid!r} 字号不一致（多数为 {majority_size}pt）",
                        auto_fix=True,
                        extra={"majority_size": majority_size, "style_id": sid},
                    ))

    return issues


def _check_d09(doc, standards: Standards) -> list[Issue]:
    """
    D09: Footer 内容一致性检查。
    - 奇偶页/首页 footer 文本不一致时报 warning
    - footer 文本含关键词黑名单时报 warning（每个关键词一条）
    """
    from .footer_ops import audit_footers

    issues: list[Issue] = []
    parts = audit_footers(doc)
    if not parts:
        return issues

    # 一致性检查：去除空 footer 后，比较所有非空 footer 的文本
    non_empty = [p for p in parts if p["text"].strip()]
    if len(non_empty) > 1:
        texts = {p["text"].strip() for p in non_empty}
        if len(texts) > 1:
            summary = "; ".join(
                f'{p["footer_type"]}={p["text"][:40]!r}' for p in non_empty
            )
            issues.append(Issue(
                rule="D09", severity="warning", para=-1, run=None,
                detail=f"footer 内容不一致（共 {len(non_empty)} 个非空 footer）: {summary}",
                auto_fix=False,
            ))

    # 关键词黑名单检查
    if standards.footer_blacklist:
        all_text = " ".join(p["text"] for p in parts)
        for kw in standards.footer_blacklist:
            if kw in all_text:
                offenders = [p["footer_type"] for p in parts if kw in p["text"]]
                issues.append(Issue(
                    rule="D09", severity="warning", para=-1, run=None,
                    detail=f"footer 含黑名单关键词 {kw!r}（出现于: {offenders}）",
                    auto_fix=False,
                ))

    return issues


# --------------------------------------------------------------------------- #
# 主检查入口                                                                    #
# --------------------------------------------------------------------------- #

ALL_RULES = ["D01", "D02", "D03", "D04", "D05", "D06", "D07", "D08", "D09"]
AUTO_FIX_RULES = ["D01", "D02", "D04", "D05", "D07", "D08"]


def check(
    doc,
    standards: Standards,
    rules: list[str] | None = None,
    para_range: tuple[int, int] | None = None,
) -> CheckResult:
    """
    扫描文档，返回 CheckResult。

    Args:
        doc:        python-docx Document
        standards:  格式标准（由调用方传入）
        rules:      要运行的规则列表，默认全部
        para_range: 扫描范围 (start, end)，默认全文
    """
    if rules is None:
        rules = ALL_RULES

    # 推断标准字体（如果 standards.font 未指定）
    inferred_font = standards.font or _infer_standard_font(doc)
    effective_standards = Standards(
        font=inferred_font,
        ascii_font=standards.ascii_font or inferred_font,
        font_size=standards.font_size,
        toc_levels=standards.toc_levels,
        double_num_patterns=standards.double_num_patterns,
    )

    style_info = _build_style_info(doc)
    issues: list[Issue] = []

    if "D01" in rules or "D02" in rules:
        active = [r for r in ["D01", "D02"] if r in rules]
        all_d0102 = _check_d01_d02(doc, effective_standards, para_range)
        issues += [i for i in all_d0102 if i.rule in active]

    if "D03" in rules:
        issues += _check_d03(doc, style_info, para_range,
                             effective_standards.double_num_patterns)
    if "D04" in rules:
        issues += _check_d04(doc, style_info, para_range)
    if "D05" in rules:
        issues += _check_d05(doc, style_info, para_range)
    if "D06" in rules:
        issues += _check_d06(doc, style_info, para_range)
    if "D07" in rules:
        issues += _check_d07(doc)
    if "D08" in rules:
        issues += _check_d08(doc, effective_standards, style_info, para_range)
    if "D09" in rules:
        issues += _check_d09(doc, effective_standards)

    return CheckResult(issues=issues, inferred_font=inferred_font)


# --------------------------------------------------------------------------- #
# 修复函数                                                                      #
# --------------------------------------------------------------------------- #

def _set_run_fonts(run_el, ea: str, ascii_font: str | None = None):
    """设置 run 的 rFonts（eastAsia + ascii + hAnsi）。"""
    rPr = run_el.find(qn("w:rPr"))
    if rPr is None:
        rPr = etree.SubElement(run_el, qn("w:rPr"))
        run_el.insert(0, rPr)
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = etree.SubElement(rPr, qn("w:rFonts"))
    rFonts.set(qn("w:eastAsia"), ea)
    if ascii_font:
        rFonts.set(qn("w:ascii"), ascii_font)
        rFonts.set(qn("w:hAnsi"), ascii_font)


def fix(
    doc,
    check_result: CheckResult,
    standards: Standards,
    rules: list[str] | None = None,
    exclude_range: tuple[int, int] | None = None,
    dry_run: bool = False,
) -> FixResult:
    """
    对 check_result 中 auto_fix=True 的问题执行修复。

    Args:
        doc:           python-docx Document（会直接修改，调用方负责 save）
        check_result:  check() 的返回值
        standards:     格式标准（需与 check 阶段一致）
        rules:         只修复哪些规则，默认全部 auto_fix 规则
        exclude_range: 排除段落范围 (lo, hi)（含两端），如前言/目录区
        dry_run:       True = 只打印，不实际修改
    """
    if rules is None:
        rules = AUTO_FIX_RULES

    effective_font = standards.font or check_result.inferred_font
    ascii_font = standards.ascii_font or effective_font

    paras = doc.paragraphs
    log: list[str] = []
    fixed = skipped = 0

    def _excluded(para_idx: int) -> bool:
        if exclude_range is None:
            return False
        return exclude_range[0] <= para_idx <= exclude_range[1]

    for issue in check_result.issues:
        if issue.rule not in rules:
            skipped += 1
            continue
        if not issue.auto_fix:
            skipped += 1
            continue
        if _excluded(issue.para):
            skipped += 1
            log.append(f"  SKIP (excluded range) para={issue.para} {issue.rule}")
            continue
        if issue.para >= len(paras):
            skipped += 1
            continue

        para = paras[issue.para]

        # ── D01/D02: 修复 run 字体 ──────────────────────────────────────── #
        if issue.rule in ("D01", "D02"):
            if not effective_font:
                skipped += 1
                log.append(f"  SKIP (no font standard) {issue.rule} para={issue.para}")
                continue
            runs = list(_iter_para_runs(para))
            if issue.run is not None and issue.run < len(runs):
                _, run_el = runs[issue.run]
                if not dry_run:
                    _set_run_fonts(run_el, effective_font, ascii_font)
                log.append(f"  FIX {issue.rule} para={issue.para} run={issue.run} → {effective_font!r}")
                fixed += 1
            else:
                skipped += 1

        # ── D04: 添加显式 outlineLvl=9（正文级别） ──────────────────────── #
        elif issue.rule == "D04":
            if not dry_run:
                pPr = para._element.get_or_add_pPr()
                existing = pPr.find(qn("w:outlineLvl"))
                if existing is not None:
                    pPr.remove(existing)
                el = etree.SubElement(pPr, qn("w:outlineLvl"))
                el.set(qn("w:val"), "9")
            log.append(f"  FIX D04 para={issue.para} → outlineLvl=9")
            fixed += 1

        # ── D05: 从兄弟段落复制 numPr ───────────────────────────────────── #
        elif issue.rule == "D05":
            ref_numpr = issue.extra.get("ref_numpr")
            if not ref_numpr:
                skipped += 1
                continue
            if not dry_run:
                pPr = para._element.get_or_add_pPr()
                existing_np = pPr.find(qn("w:numPr"))
                if existing_np is not None:
                    pPr.remove(existing_np)
                np_el = etree.SubElement(pPr, qn("w:numPr"))
                ni_el = etree.SubElement(np_el, qn("w:ilvl"))
                ni_el.set(qn("w:val"), ref_numpr["ilvl"])
                id_el = etree.SubElement(np_el, qn("w:numId"))
                id_el.set(qn("w:val"), ref_numpr["numId"])
            log.append(f"  FIX D05 para={issue.para} → numId={ref_numpr['numId']} ilvl={ref_numpr['ilvl']}")
            fixed += 1

        # ── D07: 删除 TOC \\u 开关 ───────────────────────────────────────── #
        elif issue.rule == "D07":
            for run_el in para._element.iter(qn("w:r")):
                instr = run_el.find(qn("w:instrText"))
                if instr is None or not (instr.text or "").strip().startswith("TOC"):
                    continue
                new_text = re.sub(r"\s*\\u\b", "", instr.text or "").rstrip()
                if not dry_run:
                    instr.text = new_text
                log.append(f"  FIX D07 para={issue.para} TOC instrText \\u removed")
                fixed += 1
                break
            else:
                skipped += 1

        # ── D08: 统一标题字体/字号 ───────────────────────────────────────── #
        elif issue.rule == "D08":
            majority_font = issue.extra.get("majority_font")
            majority_size = issue.extra.get("majority_size")
            for _, run_el in _iter_para_runs(para):
                rPr = run_el.find(qn("w:rPr"))
                if rPr is None and majority_font:
                    if not dry_run:
                        _set_run_fonts(run_el, majority_font)
                    continue
                if rPr is None:
                    continue
                if majority_font:
                    rFonts = rPr.find(qn("w:rFonts"))
                    if rFonts is None:
                        rFonts = etree.SubElement(rPr, qn("w:rFonts"))
                    if not dry_run:
                        rFonts.set(qn("w:eastAsia"), majority_font)
                if majority_size:
                    sz_val = str(int(majority_size * 2))
                    for sz_tag in (qn("w:sz"), qn("w:szCs")):
                        sz_el = rPr.find(sz_tag)
                        if sz_el is None:
                            sz_el = etree.SubElement(rPr, sz_tag)
                        if not dry_run:
                            sz_el.set(qn("w:val"), sz_val)
            log.append(f"  FIX D08 para={issue.para} → font={majority_font} size={majority_size}pt")
            fixed += 1

        else:
            skipped += 1

    return FixResult(fixed=fixed, skipped=skipped, dry_run=dry_run, log=log)
