"""
cli.py — lex_docx 命令行入口

用法：
    python -m lex_docx <command> [options]
    # 或软链后：
    lex_docx <command> [options]

Commands:
  ── 检查 ──────────────────────────────────────────────────────────────────────
    lint                检查 DOCX 格式（输出 JSON 或文本；支持 Profile+Selector 模式）

  ── 数据填充 ───────────────────────────────────────────────────────────────────
    extract             提取表格数据（输出 JSON）
    fill-table          按列映射填充表格
    fill-kv             填充 KV 表（基本信息类）

  ── 表格操作 ───────────────────────────────────────────────────────────────────
    format-table        统一表格格式（底色/边框/列宽/对齐）
    copy-table          跨文档表格复制（含格式）
    table-inspect       读取表格完整格式信息（底色/边框/列宽/字体/风格检测）
    table-format-brush  表格格式刷（从参考表格复制格式到目标表格）

  ── 段落/Track Changes ─────────────────────────────────────────────────────────
    tc-insert           段落级 TC INS（在指定段落插入文字）
    tc-delete           段落级 TC DEL（将指定段落标记为删除）
    highlight           批量高亮段落范围
    format-brush        格式刷（从参考段落复制格式到目标段落）
    set-outline-level   设置段落大纲级别（w:outlineLvl，独立于 Heading 样式）
    para-query          全文格式检索（按字体/样式/大纲级别/字号/粗斜体等过滤段落）

  ── 文档维护 ───────────────────────────────────────────────────────────────────
    cleanup             清理空段落 / 孤儿编号
    bold-terms          加粗定义术语
    doctor check        格式诊断（字体/编号/大纲/样式引用/TOC 开关，只读）
    doctor fix          自动修复（D01/D02/D04/D05/D07/D08，支持 --dry-run）
    inject              读取 JSON 计划文件一键执行注入
"""
from __future__ import annotations

import argparse
import json
import sys
from pathlib import Path


# =========================================================================== #
# helpers                                                                      #
# =========================================================================== #

def _load_doc(path: str):
    from docx import Document
    return Document(path)


def _save_doc(doc, path: str):
    doc.save(path)
    print(f"saved → {path}", file=sys.stderr)


def _load_cfg(cfg_path: str | None):
    from lex_docx import DocConfig
    if not cfg_path:
        return DocConfig()
    data = json.loads(Path(cfg_path).read_text(encoding="utf-8"))
    return DocConfig(**data)


def _load_json(path: str):
    return json.loads(Path(path).read_text(encoding="utf-8"))


def _out(data, fmt: str = "json"):
    if fmt == "json":
        print(json.dumps(data, ensure_ascii=False, indent=2))
    else:
        if isinstance(data, list):
            for item in data:
                print(item)
        else:
            print(data)


# =========================================================================== #
# commands                                                                     #
# =========================================================================== #

def cmd_lint(args):
    """
    lex_docx lint report.docx [--cfg config.json] [--rules rule1,rule2] [--fmt text|json]
    lex_docx lint report.docx --lint-cfg lint-config.json [--profile dd_report_draft]
    """
    from lex_docx import lint

    cfg        = _load_cfg(args.cfg) if args.cfg else None
    rules      = args.rules.split(",") if args.rules else None
    lint_cfg   = args.lint_cfg or None
    profile    = args.profile or None

    results = lint.check(
        args.docx,
        config=cfg,
        rules=rules,
        lint_cfg=lint_cfg,
        profile=profile,
    )

    # ── lint_cfg 模式：gate 判定 + 增强输出 ──────────────────────────────── #
    if lint_cfg:
        from lex_docx import lint_config as lc
        raw_cfg = lc.load_file(lint_cfg) if not isinstance(lint_cfg, dict) else lint_cfg
        resolved = lc.resolve(raw_cfg, profile_name=profile, doc_path=args.docx)
        gate_result = lc.gate_check(results, resolved.gate)

        if args.fmt == "json":
            _out({
                "profile":  resolved.name,
                "gate":     gate_result["gate"],
                "summary":  gate_result["summary"],
                "fail_reasons": gate_result["fail_reasons"],
                "results": [{
                    "rule":      r.rule,
                    "severity":  r.severity,
                    "passed":    r.passed,
                    "detail":    r.detail,
                    "locations": r.locations,
                } for r in results],
            })
        else:
            gate_icon = "✅ PASS" if gate_result["gate"] == "PASS" else "❌ FAIL"
            print(f"Profile: {resolved.name}  Gate: {gate_icon}")
            s = gate_result["summary"]
            print(f"Summary: error={s['error']} warn={s['warn']} info={s.get('info',0)}")
            if gate_result["fail_reasons"]:
                for reason in gate_result["fail_reasons"]:
                    print(f"  ⛔ {reason}")
            print()
            for r in results:
                sev_tag = {"error": "🔴", "warn": "🟡", "info": "🔵"}.get(r.severity, "⚪")
                icon = "✅" if r.passed else f"❌{sev_tag}"
                print(f"{icon} {r.rule}: {r.detail}")
                for loc in r.locations[:5]:
                    print(f"    → {loc.get('context', loc)}")
                if len(r.locations) > 5:
                    print(f"    … 共 {len(r.locations)} 处")
            sys.exit(0 if gate_result["gate"] == "PASS" else 1)
        return

    # ── 经典模式（向后兼容）──────────────────────────────────────────────── #
    if args.fmt == "json":
        _out([{
            "rule":      r.rule,
            "passed":    r.passed,
            "detail":    r.detail,
            "locations": r.locations,
        } for r in results])
    else:
        any_fail = False
        for r in results:
            icon = "✅" if r.passed else "❌"
            print(f"{icon} {r.rule}: {r.detail}")
            for loc in r.locations:
                print(f"    → {loc}")
            if not r.passed:
                any_fail = True
        sys.exit(1 if any_fail else 0)


def cmd_extract(args):
    """
    lex_docx extract source.docx --table 3 [--near "股东情况"] [--fmt json|csv]
    """
    from lex_docx import table_ops
    kwargs = {}
    if args.near:
        kwargs["near_text"] = args.near
    elif args.table is not None:
        kwargs["table_index"] = args.table
    else:
        print("error: --table or --near required", file=sys.stderr)
        sys.exit(1)

    data = table_ops.extract_table(args.docx, output="list_of_dicts", **kwargs)

    if args.fmt == "csv":
        import csv, io
        buf = io.StringIO()
        if data:
            w = csv.DictWriter(buf, fieldnames=data[0].keys())
            w.writeheader()
            w.writerows(data)
        print(buf.getvalue(), end="")
    else:
        _out(data)


def cmd_fill_table(args):
    """
    lex_docx fill-table report.docx --table 12 --data data.json
                [--map map.json] [--cfg config.json] [--out output.docx]
    """
    from lex_docx import table_ops
    doc = _load_doc(args.docx)
    data = _load_json(args.data)
    column_mapping = _load_json(args.map) if args.map else None
    cfg = _load_cfg(args.cfg)

    if args.auto_del:
        # auto delete extra rows first
        table_ops.adjust_rows(doc, args.table,
                               target_data_rows=len(data), cfg=cfg)

    filled = table_ops.fill_table(doc, args.table, data,
                                   column_mapping=column_mapping, cfg=cfg)
    _save_doc(doc, args.out or args.docx)
    _out({"filled_rows": filled})


def cmd_fill_kv(args):
    """
    lex_docx fill-kv report.docx --table 8 --data data.json
               [--key-cols 0,2] [--cfg config.json] [--out output.docx]
    """
    from lex_docx import table_ops
    doc = _load_doc(args.docx)
    data = _load_json(args.data)
    cfg = _load_cfg(args.cfg)

    key_columns = None
    if args.key_cols:
        key_columns = [int(x) for x in args.key_cols.split(",")]

    filled = table_ops.fill_kv_table(
        doc, args.table, data,
        key_columns=key_columns,
        cfg=cfg,
    )
    _save_doc(doc, args.out or args.docx)
    _out({"filled_cells": filled})


def cmd_format_table(args):
    """
    lex_docx format-table report.docx --table 12
               [--shading D9E2F3] [--borders single] [--cfg config.json] [--out output.docx]
    """
    from lex_docx import table_ops
    doc = _load_doc(args.docx)
    cfg = _load_cfg(args.cfg)

    kwargs = {}
    if args.shading:
        kwargs["header_shading"] = args.shading
    if args.borders:
        kwargs["borders"] = args.borders
    if args.widths:
        kwargs["column_widths"] = [int(x) for x in args.widths.split(",")]
    if args.align:
        kwargs["column_alignments"] = args.align.split(",")

    table_ops.format_table(doc, args.table, cfg=cfg, **kwargs)
    _save_doc(doc, args.out or args.docx)
    _out({"status": "ok"})


def cmd_cleanup(args):
    """
    lex_docx cleanup report.docx [--range 0,200] [--mode report|fix]
               [--keep-styles "Heading 1,Heading 2"] [--cfg config.json] [--out output.docx]
    """
    from lex_docx import cleanup
    cfg = _load_cfg(args.cfg)
    para_range = None
    if args.range:
        a, b = args.range.split(",")
        para_range = (int(a), int(b))
    keep_styles = args.keep_styles.split(",") if args.keep_styles else None

    doc = _load_doc(args.docx)
    as_tc_del = (args.mode != "delete")   # default: TC DEL; "delete" = hard remove

    result = cleanup.cleanup_all(
        doc,
        as_tc_del=as_tc_del,
        para_range=para_range,
        cfg=cfg,
        keep_styles=keep_styles,
    )

    if args.mode != "report":
        _save_doc(doc, args.out or args.docx)

    _out(result)


def cmd_bold_terms(args):
    """
    lex_docx bold-terms report.docx --para 39 [--out output.docx]
    lex_docx bold-terms report.docx --scan [--range 0,100]
    """
    from lex_docx import defined_terms
    doc = _load_doc(args.docx)

    if args.scan:
        para_range = None
        if args.range:
            a, b = args.range.split(",")
            para_range = (int(a), int(b))
        results = defined_terms.scan_terms(doc, para_range=para_range)
        _out(results)
        return

    if args.para is None:
        print("error: --para required (or use --scan)", file=sys.stderr)
        sys.exit(1)

    terms = defined_terms.auto_bold(doc, paragraph_index=args.para)
    _save_doc(doc, args.out or args.docx)
    _out({"bolded": terms})


def cmd_copy_table(args):
    """
    lex_docx copy-table src.docx [--src-table N | --src-near TEXT] dst.docx
               --dst-pos after_para:N|replace_table:N [--out out.docx]
    """
    from lex_docx import table_ops
    from docx import Document

    cfg = _load_cfg(args.cfg) if hasattr(args, "cfg") and args.cfg else None
    dst_doc = _load_doc(args.dst_docx)

    transform = {}
    if hasattr(args, "cols") and args.cols:
        transform["columns"] = [int(c) for c in args.cols.split(",")]
    if hasattr(args, "max_rows") and args.max_rows:
        transform["max_rows"] = args.max_rows
    if hasattr(args, "rename") and args.rename:
        import json
        transform["rename_headers"] = json.loads(args.rename)

    kwargs: dict = {}
    if hasattr(args, "src_table") and args.src_table is not None:
        kwargs["src_table_index"] = args.src_table
    if hasattr(args, "src_near") and args.src_near:
        kwargs["src_near_text"] = args.src_near

    table_ops.copy_table(
        src_doc=args.src_docx,
        dst_doc=dst_doc,
        dst_position=args.dst_pos,
        transform=transform or None,
        cfg=cfg,
        **kwargs,
    )
    _save_doc(dst_doc, args.out or args.dst_docx)
    print(f"ok: table copied → {args.out or args.dst_docx}")


def cmd_table_inspect(args):
    """
    lex_docx table-inspect report.docx --table 5 [--fmt json|text]
    """
    from lex_docx import table_ops
    kwargs = {}
    if args.near:
        kwargs["near_text"] = args.near
    else:
        kwargs["table_index"] = args.table
    result = table_ops.inspect_table(args.docx, **kwargs)
    if args.fmt == "text":
        t = result
        print(f"Table {t['table_index']}: {t['rows']}行 × {t['cols']}列  [{t['detected_style']}]")
        print(f"  列宽(dxa): {t['col_widths_dxa']}")
        print(f"  列对齐:    {t['col_aligns']}")
        print(f"  边框: {t['borders']}")
        print(f"  标题行: {t['header_row']}")
        print(f"  数据行: {t['data_rows']}")
    else:
        _out(result)


def cmd_table_format_brush(args):
    """
    lex_docx table-format-brush report.docx --ref-table 5 --target-table 12 --out out.docx
    lex_docx table-format-brush template.docx --ref-table 3 report.docx --target-table 12 --out out.docx
    """
    from lex_docx import table_ops

    # 单文档 vs 跨文档
    if args.target_docx:
        ref_src  = args.docx
        dst_doc  = _load_doc(args.target_docx)
        out_path = args.out or args.target_docx
    else:
        dst_doc  = _load_doc(args.docx)
        ref_src  = dst_doc
        out_path = args.out or args.docx

    copy = args.copy.split(",") if args.copy else None
    result = table_ops.table_format_brush(
        ref_src, args.ref_table, dst_doc, args.target_table, copy=copy
    )
    _save_doc(dst_doc, out_path)
    _out({"ok": True, **result})


def cmd_tc_insert(args):
    """
    lex_docx tc-insert report.docx --para 180 --text "新增文字"
               [--pos end|start|N] [--bold] [--italic] [--highlight yellow]
               [--inherit-rpr true|style|auto] [--cfg config.json] [--out out.docx]
    """
    from lex_docx.tc_utils import tc_ins_text, next_tc_id
    cfg = _load_cfg(args.cfg)
    doc = _load_doc(args.docx)

    author = cfg.author if cfg else "JT"
    tc_id  = next_tc_id(doc)
    para   = doc.paragraphs[args.para]

    inherit = args.inherit_rpr
    if inherit == "true":
        inherit = True

    tc_ins_text(
        para._element,
        text=args.text,
        tc_id=tc_id,
        author=author,
        position=args.pos,
        bold=args.bold,
        italic=args.italic,
        highlight=args.highlight or None,
        inherit_rPr=inherit,
    )
    _save_doc(doc, args.out or args.docx)
    _out({"ok": True, "para": args.para, "text": args.text})


def cmd_tc_delete(args):
    """
    lex_docx tc-delete report.docx --para 180
    lex_docx tc-delete report.docx --range 180,195
               [--cfg config.json] [--out out.docx]
    """
    from lex_docx.tc_utils import tc_del_paragraph, next_tc_id
    cfg    = _load_cfg(args.cfg)
    doc    = _load_doc(args.docx)
    author = cfg.author if cfg else "JT"
    tc_id  = next_tc_id(doc)

    if args.range:
        a, b = args.range.split(",")
        indices = list(range(int(a), int(b) + 1))
    elif args.para is not None:
        indices = [args.para]
    else:
        print("error: --para or --range required", file=sys.stderr)
        sys.exit(1)

    deleted = []
    for idx in indices:
        if idx >= len(doc.paragraphs):
            continue
        tc_del_paragraph(doc.paragraphs[idx]._element, tc_id, author)
        deleted.append(idx)
        tc_id += 1

    _save_doc(doc, args.out or args.docx)
    _out({"ok": True, "deleted": deleted})


def cmd_highlight(args):
    """
    lex_docx highlight report.docx --range 180,195
    lex_docx highlight report.docx --para 180
               [--color yellow] [--out out.docx]
    """
    from lxml import etree
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = _load_doc(args.docx)
    color = args.color or "yellow"

    if args.range:
        a, b = args.range.split(",")
        indices = list(range(int(a), int(b) + 1))
    elif args.para is not None:
        indices = [args.para]
    else:
        print("error: --para or --range required", file=sys.stderr)
        sys.exit(1)

    marked = []
    for idx in indices:
        if idx >= len(doc.paragraphs):
            continue
        para_el = doc.paragraphs[idx]._element
        runs = para_el.findall(qn("w:r"))
        if not runs:
            # 段落无 run — 创建空 run 承载 highlight
            r = OxmlElement("w:r")
            rPr = OxmlElement("w:rPr")
            hl = OxmlElement("w:highlight")
            hl.set(qn("w:val"), color)
            rPr.append(hl)
            r.insert(0, rPr)
            para_el.append(r)
        else:
            for run_el in runs:
                rPr = run_el.find(qn("w:rPr"))
                if rPr is None:
                    rPr = OxmlElement("w:rPr")
                    run_el.insert(0, rPr)
                existing = rPr.find(qn("w:highlight"))
                if existing is not None:
                    rPr.remove(existing)
                hl = OxmlElement("w:highlight")
                hl.set(qn("w:val"), color)
                rPr.append(hl)
        marked.append(idx)

    _save_doc(doc, args.out or args.docx)
    _out({"ok": True, "highlighted": marked, "color": color})


def cmd_format_brush(args):
    """
    lex_docx format-brush report.docx --ref 171 --target 177,178,180
    lex_docx format-brush report.docx --ref 171 --range 175,185
               [--copy indent,spacing,style] [--out out.docx]
    """
    from lex_docx import format_brush

    doc = _load_doc(args.docx)

    if args.target:
        indices = [int(x) for x in args.target.split(",")]
    elif args.range:
        a, b = args.range.split(",")
        indices = list(range(int(a), int(b) + 1))
    else:
        print("error: --target or --range required", file=sys.stderr)
        sys.exit(1)

    copy = args.copy.split(",") if args.copy else None

    modified = format_brush.apply(
        doc,
        target_indices=indices,
        reference_index=args.ref,
        copy=copy,
    )
    _save_doc(doc, args.out or args.docx)
    _out({"ok": True, "modified": modified})


def cmd_set_outline_level(args):
    """
    lex_docx set-outline-level report.docx --para 5 --level 2
    lex_docx set-outline-level report.docx --range 10,20 --level 2
    lex_docx set-outline-level report.docx --style "自定义标题" --level 1
    lex_docx set-outline-level report.docx --range 0,200 --style "自定义标题" --level 1
    lex_docx set-outline-level report.docx --para 5 --level none  # 清除大纲级别
    """
    from lex_docx import format_brush

    doc = _load_doc(args.docx)
    paras = doc.paragraphs

    # 确定目标段落索引
    if args.para is not None:
        indices: list[int] = [args.para]
    elif args.range:
        lo, hi = [int(x.strip()) for x in args.range.split(",", 1)]
        indices = list(range(lo, hi + 1))
    else:
        indices = list(range(len(paras)))

    # 按 style 过滤
    if args.style:
        indices = [i for i in indices
                   if i < len(paras) and paras[i].style and paras[i].style.name == args.style]

    # 解析 level：整数 1-9 或 "none"/0 表示清除
    if args.level.lower() == "none":
        level = None
    else:
        level = int(args.level)

    modified = format_brush.set_outline_level(doc, indices, level)
    _save_doc(doc, args.out or args.docx)
    _out({"modified": len(modified), "indices": modified})


def cmd_doctor(args):
    """
    lex_docx doctor check report.docx --font 楷体 --font-size 12
    lex_docx doctor fix   report.docx --font 楷体 --rules D01,D02,D04 --dry-run
    """
    from lex_docx import doctor as dr

    standards = dr.Standards(
        font=args.font or None,
        ascii_font=args.ascii_font or None,
        font_size=float(args.font_size) if args.font_size else None,
        toc_levels=tuple(int(x) for x in args.toc_levels.split("-")) if args.toc_levels else (1, 3),
    )
    rules = [r.strip().upper() for r in args.rules.split(",")] if args.rules else None

    para_range = None
    if args.range:
        lo, hi = [int(x.strip()) for x in args.range.split(",", 1)]
        para_range = (lo, hi + 1)

    doc = _load_doc(args.docx)

    if args.doctor_cmd == "check":
        result = dr.check(doc, standards, rules=rules, para_range=para_range)
        _out(result.to_dict())

    elif args.doctor_cmd == "fix":
        # check 先跑一遍
        check_result = dr.check(doc, standards, rules=rules, para_range=para_range)

        exclude_range = None
        if args.exclude_range:
            lo, hi = [int(x.strip()) for x in args.exclude_range.split(",", 1)]
            exclude_range = (lo, hi)

        if args.backup and not args.dry_run:
            import shutil
            shutil.copy2(args.docx, args.docx + ".bak")

        fix_result = dr.fix(
            doc, check_result, standards,
            rules=rules,
            exclude_range=exclude_range,
            dry_run=args.dry_run,
        )

        if not args.dry_run:
            _save_doc(doc, args.out or args.docx)

        _out({
            "dry_run":   fix_result.dry_run,
            "fixed":     fix_result.fixed,
            "skipped":   fix_result.skipped,
            "log":       fix_result.log,
            "check_summary": check_result.summary(),
        })


def cmd_para_query(args):
    """
    lex_docx para-query report.docx --font "仿宋"
    lex_docx para-query report.docx --outline-level 1,2
    lex_docx para-query report.docx --style "Heading 1" "Heading 2"
    lex_docx para-query report.docx --font "仿宋" --font-size 12
    lex_docx para-query report.docx --bold
    lex_docx para-query report.docx --range 0,200 --font "仿宋"
    """
    from lex_docx import para_query

    doc = _load_doc(args.docx)

    outline_level = None
    if args.outline_level:
        outline_level = [int(x.strip()) for x in args.outline_level.split(",")]

    para_range = None
    if args.range:
        lo, hi = [int(x.strip()) for x in args.range.split(",", 1)]
        para_range = (lo, hi + 1)

    bold = None
    if args.bold:
        bold = True
    elif args.no_bold:
        bold = False

    italic = None
    if args.italic:
        italic = True
    elif args.no_italic:
        italic = False

    results = para_query.query(
        doc,
        style=args.style or None,
        font=args.font,
        font_size=float(args.font_size) if args.font_size else None,
        outline_level=outline_level,
        bold=bold,
        italic=italic,
        color=args.color,
        para_range=para_range,
        text_preview_len=args.preview_len,
    )

    if args.fmt == "text":
        for r in results:
            ol = f"大纲{r['outline_level']}级" if r["outline_level"] else ""
            font_str = "/".join(r["font_eastasia"] or r["font_ascii"] or [])
            sz_str = "/".join(str(s) for s in r["font_size"]) + "pt" if r["font_size"] else ""
            flags = " ".join(f for f, v in [("粗", r["bold"]), ("斜", r["italic"])] if v)
            meta = "  ".join(x for x in [r["style"], ol, font_str, sz_str, flags] if x)
            print(f"[{r['index']:>4}] {meta}")
            print(f"       {r['text']}")
    else:
        _out(results)


def cmd_inject(args):
    """
    lex_docx inject plan.json [--cfg config.json] [--out out.docx]

    plan.json 结构：
    {
      "doc_path": "report.docx",
      "out_path": "report_out.docx",    // 可被 --out 覆盖
      "target_range": [200, 300],       // 可选
      "tables": [
        {"table_index": 8, "data": {...}, "mode": "kv", "key_columns": [0, 2]},
        {"table_index": 12, "data": [...], "mode": "rows", "auto_adjust": true}
      ],
      "jt_notes": {"180": "待确认", "治理结构": "待核实"},
      "auto_cleanup": true,
      "run_lint": true
    }
    """
    from lex_docx import inject_engine

    raw = _load_json(args.plan)
    cfg = _load_cfg(args.cfg)

    # out path: CLI --out 优先于 JSON 中的 out_path
    if args.out:
        raw["out_path"] = args.out

    # jt_notes key 从 JSON string 转回 int（JSON key 只能是 str）
    jt_notes_raw = raw.pop("jt_notes", {})
    jt_notes: dict = {}
    for k, v in jt_notes_raw.items():
        try:
            jt_notes[int(k)] = v
        except ValueError:
            jt_notes[k] = v

    tables = [inject_engine.TableFill(**t) for t in raw.pop("tables", [])]

    plan = inject_engine.InjectPlan(tables=tables, jt_notes=jt_notes, **raw)
    result = inject_engine.execute(plan, cfg)
    _out({
        "summary": result.summary(),
        "tables":  result.tables,
        "notes":   result.notes,
        "cleanup": result.cleanup,
    })


# =========================================================================== #
# main                                                                         #
# =========================================================================== #

def main():
    parser = argparse.ArgumentParser(
        prog="lex_docx",
        usage="lex_docx <command> [options]  (lex_docx --help 查看全部命令)",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        description="""\
lex_docx — DOCX 自动化工具库 CLI

Commands:
  ── 检查 / 诊断 ──────────────────────────────────────────────────────────────
  lint                检查 DOCX 内容格式（拼写/表格/标注，支持 Profile+Selector）
  doctor check        格式结构诊断（字体/大纲/编号/样式引用/TOC 开关，只读）
  doctor fix          自动修复诊断结果（支持 --dry-run / --backup）

  ── 数据填充 ────────────────────────────────────────────────────────────────
  extract             提取表格数据（输出 JSON 或 CSV）
  fill-table          按列映射批量填充表格行
  fill-kv             填充 KV 表（基本信息类二列/四列布局）

  ── 表格操作 ────────────────────────────────────────────────────────────────
  format-table        统一表格格式（底色/边框/列宽/对齐）
  copy-table          跨文档表格复制（含完整格式）
  table-inspect       读取表格完整格式信息（底色/边框/列宽/字体/风格检测）
  table-format-brush  表格格式刷（从参考表格复制格式到目标表格）

  ── 段落 / Track Changes ────────────────────────────────────────────────────
  tc-insert           段落级 TC INS（在指定段落插入带标记文字）
  tc-delete           段落级 TC DEL（将指定段落标记为删除）
  highlight           批量高亮段落范围
  format-brush        段落格式刷（复制缩进/间距/样式/大纲级别等）
  set-outline-level   设置段落大纲级别（w:outlineLvl，独立于 Heading 样式）
  para-query          全文格式检索（按字体/样式/大纲级别/字号/粗斜体等过滤）

  ── 文档维护 ────────────────────────────────────────────────────────────────
  cleanup             清理空段落 / 孤儿编号
  bold-terms          加粗定义术语首次出现位置
  inject              读取 JSON 计划文件一键批量注入

每个子命令均支持 -h / --help 查看详细参数。
""",
    )
    sub = parser.add_subparsers(dest="command", required=True,
                                metavar="<command>",
                                help=argparse.SUPPRESS)

    # ── lint ──────────────────────────────────────────────────────────────── #
    p = sub.add_parser("lint", help="检查 DOCX 内容格式（支持 Profile+Selector 模式）")
    p.add_argument("docx")
    p.add_argument("--cfg", help="DocConfig JSON 文件路径（经典模式）")
    p.add_argument("--rules", help="逗号分隔的规则名，默认全部")
    p.add_argument("--fmt", choices=["text", "json"], default="text")
    p.add_argument("--lint-cfg", dest="lint_cfg",
                   help="Lint Config JSON 路径（Profile + Selector 模式）")
    p.add_argument("--profile", help="指定 profile 名；不指定则按 selectors 自动匹配")

    # ── extract ───────────────────────────────────────────────────────────── #
    p = sub.add_parser("extract", help="提取表格数据")
    p.add_argument("docx")
    p.add_argument("--table", type=int, help="表格索引")
    p.add_argument("--near", help="临近文字定位")
    p.add_argument("--fmt", choices=["json", "csv"], default="json")

    # ── fill-table ────────────────────────────────────────────────────────── #
    p = sub.add_parser("fill-table", help="按列映射填充表格")
    p.add_argument("docx")
    p.add_argument("--table", type=int, required=True)
    p.add_argument("--data", required=True, help="List[Dict] JSON 文件")
    p.add_argument("--map", help="列映射 Dict JSON 文件")
    p.add_argument("--auto-del", action="store_true", help="自动 TC DEL 多余行")
    p.add_argument("--cfg")
    p.add_argument("--out", help="输出路径，默认覆盖原文件")

    # ── fill-kv ───────────────────────────────────────────────────────────── #
    p = sub.add_parser("fill-kv", help="填充 KV 表")
    p.add_argument("docx")
    p.add_argument("--table", type=int, required=True)
    p.add_argument("--data", required=True, help="Dict JSON 文件")
    p.add_argument("--key-cols", help="多列 key 索引，如 '0,2'（四列布局）")
    p.add_argument("--cfg")
    p.add_argument("--out")

    # ── format-table ──────────────────────────────────────────────────────── #
    p = sub.add_parser("format-table", help="统一表格格式")
    p.add_argument("docx")
    p.add_argument("--table", type=int, required=True)
    p.add_argument("--shading", help="标题行底色十六进制，如 D9E2F3")
    p.add_argument("--borders", choices=["single", "none"])
    p.add_argument("--widths", help="列宽 dxa，逗号分隔，如 '800,4000,2000'")
    p.add_argument("--align", help="列对齐，逗号分隔，如 'center,left,right'")
    p.add_argument("--cfg")
    p.add_argument("--out")

    # ── cleanup ───────────────────────────────────────────────────────────── #
    p = sub.add_parser("cleanup", help="清理空段落 / 孤儿编号")
    p.add_argument("docx")
    p.add_argument("--range", help="段落范围，如 '0,200'")
    p.add_argument("--mode", choices=["report", "tc-del", "delete"],
                   default="tc-del",
                   help="report=只报告 | tc-del=TC DEL标记（默认）| delete=直接删除")
    p.add_argument("--keep-styles", help="保留这些 style 的空段落，逗号分隔")
    p.add_argument("--cfg")
    p.add_argument("--out")

    # ── bold-terms ────────────────────────────────────────────────────────── #
    p = sub.add_parser("bold-terms", help="加粗定义术语")
    p.add_argument("docx")
    p.add_argument("--para", type=int, help="目标段落索引")
    p.add_argument("--scan", action="store_true", help="扫描全文，只查不改")
    p.add_argument("--range", help="scan 范围，如 '0,100'")
    p.add_argument("--out")

    # ── copy-table ────────────────────────────────────────────────────────── #
    p = sub.add_parser("copy-table", help="跨文档表格复制（含格式）")
    p.add_argument("src_docx", help="源文档路径")
    p.add_argument("dst_docx", help="目标文档路径")
    p.add_argument("--dst-pos", required=True,
                   help="插入位置：after_para:N 或 replace_table:N")
    p.add_argument("--src-table", type=int, help="源表格序号（0-based）")
    p.add_argument("--src-near", help="按临近文字定位源表格")
    p.add_argument("--cols", help="保留列号，逗号分隔，如 '0,1,2,3'")
    p.add_argument("--max-rows", type=int, help="最多保留数据行数")
    p.add_argument("--rename", help='重命名表头 JSON，如 \'{"变更时间":"日期"}\'')
    p.add_argument("--cfg")
    p.add_argument("--out", help="输出路径，默认覆盖 dst_docx")

    # ── table-inspect ─────────────────────────────────────────────────────── #
    p = sub.add_parser("table-inspect", help="读取表格完整格式信息（底色/边框/列宽/字体）")
    p.add_argument("docx")
    p.add_argument("--table", type=int, help="表格序号（0-based）")
    p.add_argument("--near",  help="按临近文字定位表格")
    p.add_argument("--fmt", choices=["json", "text"], default="json")

    # ── table-format-brush ────────────────────────────────────────────────── #
    p = sub.add_parser("table-format-brush", help="表格格式刷（从参考表格复制格式到目标表格）")
    p.add_argument("docx", help="参考文档（单文档模式时也是目标文档）")
    p.add_argument("target_docx", nargs="?", help="目标文档（跨文档模式时指定）")
    p.add_argument("--ref-table",    dest="ref_table",    type=int, required=True,
                   help="参考表格序号")
    p.add_argument("--target-table", dest="target_table", type=int, required=True,
                   help="目标表格序号")
    p.add_argument("--copy", help="复制项，逗号分隔：shading,borders,col_widths,col_aligns,font,row_height")
    p.add_argument("--out")

    # ── tc-insert ─────────────────────────────────────────────────────────── #
    p = sub.add_parser("tc-insert", help="段落级 TC INS（在指定段落插入文字）")
    p.add_argument("docx")
    p.add_argument("--para", type=int, required=True, help="目标段落索引")
    p.add_argument("--text", required=True, help="插入的文字内容")
    p.add_argument("--pos", default="end", help="插入位置：end（默认）| start | N（整数）")
    p.add_argument("--bold", action="store_true")
    p.add_argument("--italic", action="store_true")
    p.add_argument("--highlight", help="高亮颜色，如 yellow")
    p.add_argument("--inherit-rpr", dest="inherit_rpr", default="true",
                   choices=["true", "style", "auto"],
                   help="rPr 继承策略：true=继承首 run（默认）| style=跟 pStyle | auto=按 style_rPr_map")
    p.add_argument("--cfg")
    p.add_argument("--out")

    # ── tc-delete ─────────────────────────────────────────────────────────── #
    p = sub.add_parser("tc-delete", help="段落级 TC DEL（将指定段落标记为删除）")
    p.add_argument("docx")
    p.add_argument("--para", type=int, help="单个段落索引")
    p.add_argument("--range", help="段落范围（含两端），如 '180,195'")
    p.add_argument("--cfg")
    p.add_argument("--out")

    # ── highlight ─────────────────────────────────────────────────────────── #
    p = sub.add_parser("highlight", help="批量标黄段落范围")
    p.add_argument("docx")
    p.add_argument("--para", type=int, help="单个段落索引")
    p.add_argument("--range", help="段落范围（含两端），如 '180,195'")
    p.add_argument("--color", default="yellow", help="高亮颜色（默认 yellow）")
    p.add_argument("--out")

    # ── format-brush ──────────────────────────────────────────────────────── #
    p = sub.add_parser("format-brush", help="格式刷（从参考段落复制格式到目标段落）")
    p.add_argument("docx")
    p.add_argument("--ref", type=int, required=True, help="参考段落索引")
    p.add_argument("--target", help="目标段落索引，逗号分隔，如 '177,178,180'")
    p.add_argument("--range", help="目标段落范围（含两端），如 '175,185'")
    p.add_argument("--copy", help="复制项，逗号分隔，如 'indent,spacing,style'（默认全部）")
    p.add_argument("--out")

    # ── set-outline-level ─────────────────────────────────────────────────── #
    p = sub.add_parser("set-outline-level",
                       help="设置段落大纲级别（w:outlineLvl，独立于 Heading 样式）")
    p.add_argument("docx")
    p.add_argument("--level", required=True,
                   help="大纲级别 1-9，或 none/0 表示清除（变为正文）")
    p.add_argument("--para",  type=int, help="单个段落索引")
    p.add_argument("--range", help="段落范围（含两端），如 '10,20'")
    p.add_argument("--style", help="按样式名过滤，如 '自定义标题'")
    p.add_argument("--out",   help="输出路径，默认覆盖原文件")

    # ── para-query ────────────────────────────────────────────────────────── #
    p = sub.add_parser("para-query",
                       help="全文格式检索（按字体/样式/大纲级别/字号/粗斜体等过滤段落）")
    p.add_argument("docx")
    p.add_argument("--style",    nargs="+", help="段落样式名（可多个，OR 匹配），如 --style 'Heading 1' '标题1'")
    p.add_argument("--font",     help="字体名（部分匹配），如 '仿宋'")
    p.add_argument("--font-size", dest="font_size",
                   help="字号（pt），如 '12'")
    p.add_argument("--outline-level", dest="outline_level",
                   help="大纲级别，逗号分隔，如 '1,2'")
    p.add_argument("--bold",     action="store_true", default=False,
                   help="只返回含粗体 run 的段落")
    p.add_argument("--no-bold",  dest="no_bold", action="store_true", default=False,
                   help="只返回不含粗体 run 的段落")
    p.add_argument("--italic",   action="store_true", default=False,
                   help="只返回含斜体 run 的段落")
    p.add_argument("--no-italic", dest="no_italic", action="store_true", default=False,
                   help="只返回不含斜体 run 的段落")
    p.add_argument("--color",    help="字体颜色十六进制，如 'FF0000'")
    p.add_argument("--range",    help="扫描范围（含两端），如 '0,200'")
    p.add_argument("--preview-len", dest="preview_len", type=int, default=60,
                   help="文本预览截断长度（默认 60）")
    p.add_argument("--fmt",      choices=["json", "text"], default="json")

    # ── doctor ────────────────────────────────────────────────────────────── #
    p = sub.add_parser("doctor", help="格式诊断（check）与自动修复（fix）")
    doctor_sub = p.add_subparsers(dest="doctor_cmd", required=True)

    _doctor_shared_args = lambda pp: (
        pp.add_argument("docx"),
        pp.add_argument("--font",        help="标准 eastAsia 字体，如 '楷体'（不指定则自动推断）"),
        pp.add_argument("--ascii-font",  dest="ascii_font",
                        help="标准 ascii/hAnsi 字体（不指定则同 --font）"),
        pp.add_argument("--font-size",   dest="font_size",
                        help="标准字号 pt，如 '12'（不指定则不检查 D02 字号）"),
        pp.add_argument("--toc-levels",  dest="toc_levels",
                        help="TOC 应收录的大纲级别范围，如 '1-3'（默认 1-3）"),
        pp.add_argument("--rules",       help="只运行指定规则，逗号分隔，如 'D01,D02,D04'"),
        pp.add_argument("--range",       help="扫描段落范围（含两端），如 '60,500'"),
    )

    p_check = doctor_sub.add_parser("check", help="诊断（只读，输出问题报告）")
    _doctor_shared_args(p_check)

    p_fix = doctor_sub.add_parser("fix", help="自动修复（默认修复所有 auto_fix 规则）")
    _doctor_shared_args(p_fix)
    p_fix.add_argument("--dry-run",       dest="dry_run", action="store_true",
                       help="只打印修复计划，不实际修改")
    p_fix.add_argument("--exclude-range", dest="exclude_range",
                       help="排除段落范围（含两端），如 '0,59'（前言/目录区）")
    p_fix.add_argument("--backup",        action="store_true",
                       help="修复前自动创建 .bak 备份")
    p_fix.add_argument("--out",           help="输出路径（默认覆盖原文件）")

    # ── inject ────────────────────────────────────────────────────────────── #
    p = sub.add_parser("inject", help="读取 JSON 计划文件一键执行注入")
    p.add_argument("plan", help="InjectPlan JSON 文件路径")
    p.add_argument("--cfg", help="DocConfig JSON 文件路径")
    p.add_argument("--out", help="输出路径（覆盖 plan 中的 out_path）")

    args = parser.parse_args()
    dispatch = {
        "lint":         cmd_lint,
        "extract":      cmd_extract,
        "fill-table":   cmd_fill_table,
        "fill-kv":      cmd_fill_kv,
        "format-table": cmd_format_table,
        "cleanup":      cmd_cleanup,
        "bold-terms":   cmd_bold_terms,
        "copy-table":   cmd_copy_table,
        "table-inspect":       cmd_table_inspect,
        "table-format-brush":  cmd_table_format_brush,
        "tc-insert":    cmd_tc_insert,
        "tc-delete":    cmd_tc_delete,
        "highlight":    cmd_highlight,
        "format-brush":       cmd_format_brush,
        "set-outline-level":  cmd_set_outline_level,
        "para-query":         cmd_para_query,
        "doctor":             cmd_doctor,
        "inject":             cmd_inject,
    }
    dispatch[args.command](args)


if __name__ == "__main__":
    main()
