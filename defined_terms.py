"""
defined_terms.py — 定义术语加粗

支持中文和英文合同/报告中的各类术语定义模式：

中文模式：
  （"xxx"）/ （"xxx"或"yyy"）
  以下简称"xxx" / 合称"xxx" / 下称"xxx"

英文模式：
  ("xxx") / ('xxx')
  hereinafter referred to as "xxx"
  hereinafter "xxx" / hereafter "xxx"
  collectively "xxx" / together "xxx"
  each a "xxx" / each an "xxx"

扩展：通过 extra_patterns 传入额外正则（每个含1个捕获组）。

只对引号内术语文字加粗（w:b + w:bCs），其余格式属性不变。
"""
from __future__ import annotations

import re
from copy import deepcopy

from docx.oxml import OxmlElement
from docx.oxml.ns import qn


# --------------------------------------------------------------------------- #
# 引号字符集                                                                    #
# --------------------------------------------------------------------------- #

# 中文 + 英文直引号 + 弯引号 + 中文书名号式引号
_OPEN_QUOTES  = r'["\u201c\u2018\u300c\u300e\u0022\u0027]'   # " " ' 「 『 " '
_CLOSE_QUOTES = r'["\u201d\u2019\u300d\u300f\u0022\u0027]'   # " " ' 」 』 " '
_ANY_QUOTE    = r'["\u201c\u2018\u300c\u300e\u201d\u2019\u300d\u300f\u0022\u0027]'
# 术语内部不含引号，中英文均允许，长度 1-60 字符
_TERM_BODY    = r'[^"\u201c\u2018\u300c\u300e\u201d\u2019\u300d\u300f\u0022\u0027]{1,60}'


# --------------------------------------------------------------------------- #
# 内置模式                                                                      #
# --------------------------------------------------------------------------- #

# 中文括号块：（...），内部含任意多个引号术语
_ZH_PAREN = re.compile(
    r'（[^（）]{1,200}）'
)

# 中文行内简称
_ZH_INLINE = re.compile(
    r'(?:以下简称|合称|下称|以下合称)' + _OPEN_QUOTES + r'(' + _TERM_BODY + r')' + _CLOSE_QUOTES
)

# 英文括号块：("xxx") / ('xxx') / (each a "xxx") / (the "xxx") / (together, the "xxx")
_EN_PAREN = re.compile(
    r'\((?:each\s+an?\s+|the\s+|together(?:,\s*the\s+)?|collectively(?:,\s*the\s+)?|each\s+)?'
    + _OPEN_QUOTES + r'(' + _TERM_BODY + r')' + _CLOSE_QUOTES + r'\)'
)

# 英文行内简称
_EN_INLINE = re.compile(
    r'(?:hereinafter(?:\s+referred\s+to\s+as)?|hereafter|collectively|together)\s+'
    + _OPEN_QUOTES + r'(' + _TERM_BODY + r')' + _CLOSE_QUOTES,
    re.IGNORECASE
)

# 提取引号内内容（用于中文括号块内遍历）
_QUOTED_INNER = re.compile(
    _OPEN_QUOTES + r'(' + _TERM_BODY + r')' + _CLOSE_QUOTES
)


# --------------------------------------------------------------------------- #
# 主接口                                                                        #
# --------------------------------------------------------------------------- #

def auto_bold(
    doc,
    paragraph_index: int,
    extra_patterns: list[str] | None = None,
    cfg=None,
) -> list[str]:
    """
    自动检测段落中所有定义术语并加粗。

    支持中英文混合段落。

    Args:
        doc:               python-docx Document
        paragraph_index:   目标段落索引
        extra_patterns:    额外正则字符串列表，每个含 1 个捕获组表示术语文字
                           例：[r'hereinafter\\s+"([^"]+)"']
        cfg:               DocConfig，其 extra_term_patterns 属性也会被合并

    Returns:
        加粗的术语字符串列表
    """
    para = doc.paragraphs[paragraph_index]
    full = _para_full_text(para)
    extra = _merge_extra_patterns(extra_patterns, cfg)
    spans = _find_term_spans(full, extra)
    for start, end in sorted(spans, reverse=True):
        _apply_bold_to_span(para, start, end)
    return [full[s:e] for s, e in spans]


def bold_terms(
    doc,
    paragraph_index: int,
    terms: list[str],
) -> list[str]:
    """
    手动指定术语列表并加粗（精确字符串匹配，不含引号）。

    Returns:
        实际找到并加粗的术语列表
    """
    para = doc.paragraphs[paragraph_index]
    full = _para_full_text(para)
    spans = []
    for term in terms:
        pos = full.find(term)
        if pos != -1:
            spans.append((pos, pos + len(term)))

    for start, end in sorted(spans, key=lambda x: x[0], reverse=True):
        _apply_bold_to_span(para, start, end)

    return [full[s:e] for s, e in spans]


def scan_terms(
    doc,
    para_range: tuple[int, int] | None = None,
    extra_patterns: list[str] | None = None,
    cfg=None,
) -> list[dict]:
    """
    扫描段落范围，返回所有检测到的定义术语（可在执行前确认）。

    Returns:
        [{"index": int, "text_preview": str, "terms": [str, ...]}, ...]
    """
    paras = doc.paragraphs
    start, end = para_range if para_range else (0, len(paras))
    extra = _merge_extra_patterns(extra_patterns, cfg)
    results = []
    for i in range(start, min(end, len(paras))):
        full = _para_full_text(paras[i])
        spans = _find_term_spans(full, extra)
        if spans:
            results.append({
                "index": i,
                "text_preview": full[:80],
                "terms": [full[s:e] for s, e in spans],
            })
    return results


# --------------------------------------------------------------------------- #
# 模式匹配                                                                      #
# --------------------------------------------------------------------------- #

def _merge_extra_patterns(extra_patterns, cfg) -> list[re.Pattern]:
    """合并 extra_patterns 参数和 cfg.extra_term_patterns，编译为 Pattern 列表。"""
    raw = list(extra_patterns or [])
    if cfg is not None and hasattr(cfg, "extra_term_patterns"):
        raw.extend(cfg.extra_term_patterns or [])
    compiled = []
    for p in raw:
        try:
            compiled.append(re.compile(p, re.IGNORECASE))
        except re.error:
            pass
    return compiled


def _find_term_spans(text: str, extra_patterns: list | None = None) -> list[tuple[int, int]]:
    """
    在 text 中找出所有定义术语的字符跨度 (start, end)。
    跨度仅指术语文字本身（不含引号和括号）。
    """
    spans = []

    # ── 中文括号块：遍历块内所有引号术语 ──────────────────────────────────── #
    for m in _ZH_PAREN.finditer(text):
        block      = m.group()
        block_off  = m.start()
        for inner in _QUOTED_INNER.finditer(block):
            spans.append((block_off + inner.start(1), block_off + inner.end(1)))

    # ── 中文行内简称 ──────────────────────────────────────────────────────── #
    for m in _ZH_INLINE.finditer(text):
        spans.append((m.start(1), m.end(1)))

    # ── 英文括号块 ────────────────────────────────────────────────────────── #
    for m in _EN_PAREN.finditer(text):
        spans.append((m.start(1), m.end(1)))

    # ── 英文行内简称 ──────────────────────────────────────────────────────── #
    for m in _EN_INLINE.finditer(text):
        spans.append((m.start(1), m.end(1)))

    # ── 自定义 extra 模式 ─────────────────────────────────────────────────── #
    if extra_patterns:
        for pat in extra_patterns:
            for m in pat.finditer(text):
                if m.lastindex and m.lastindex >= 1:
                    spans.append((m.start(1), m.end(1)))

    # 去重、排序、过滤空跨度
    return sorted({s for s in spans if s[0] < s[1]})


# --------------------------------------------------------------------------- #
# Run 拆分 & 加粗                                                               #
# --------------------------------------------------------------------------- #

def _para_full_text(para) -> str:
    return "".join(t.text or "" for t in para._element.iter(qn("w:t")))


def _get_all_runs_with_pos(para) -> list[tuple]:
    pos = 0
    result = []
    for run_el in para._element.iter(qn("w:r")):
        text = "".join(t.text or "" for t in run_el.findall(qn("w:t")))
        result.append((run_el, pos, pos + len(text)))
        pos += len(text)
    return result


def _set_xml_space(t_el, text: str):
    _XML_SPACE = "{http://www.w3.org/XML/1998/namespace}space"
    if text and (text[0] == " " or text[-1] == " "):
        t_el.set(_XML_SPACE, "preserve")
    else:
        t_el.attrib.pop(_XML_SPACE, None)


def _split_run_el(run_el, offset: int) -> tuple:
    """在 offset 处拆分 run，返回 (left_el, right_el)。"""
    t_el = run_el.find(qn("w:t"))
    if t_el is None:
        return run_el, None
    full = t_el.text or ""
    if offset <= 0:
        return None, run_el
    if offset >= len(full):
        return run_el, None

    left_text  = full[:offset]
    right_text = full[offset:]

    t_el.text = left_text
    _set_xml_space(t_el, left_text)

    right_el = deepcopy(run_el)
    right_t  = right_el.find(qn("w:t"))
    right_t.text = right_text
    _set_xml_space(right_t, right_text)
    run_el.addnext(right_el)

    return run_el, right_el


def _apply_bold_to_span(para, char_start: int, char_end: int):
    """对段落 [char_start, char_end) 内的 run 加粗，按需拆分边界 run。"""
    if char_start >= char_end:
        return

    # 先拆右边界（不影响左边界位置）
    for run_el, r_start, r_end in _get_all_runs_with_pos(para):
        if r_start < char_end < r_end:
            _split_run_el(run_el, char_end - r_start)
            break

    # 再拆左边界
    for run_el, r_start, r_end in _get_all_runs_with_pos(para):
        if r_start < char_start < r_end:
            _split_run_el(run_el, char_start - r_start)
            break

    # 加粗位于 [char_start, char_end) 内的 run
    for run_el, r_start, r_end in _get_all_runs_with_pos(para):
        if r_start >= char_start and r_end <= char_end:
            _ensure_bold(run_el)


def _ensure_bold(run_el):
    rPr = run_el.find(qn("w:rPr"))
    if rPr is None:
        rPr = OxmlElement("w:rPr")
        run_el.insert(0, rPr)
    if rPr.find(qn("w:b")) is None:
        rPr.append(OxmlElement("w:b"))
    if rPr.find(qn("w:bCs")) is None:
        rPr.append(OxmlElement("w:bCs"))
